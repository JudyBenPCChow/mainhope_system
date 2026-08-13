#!/usr/bin/env python3
"""Generate 2627 regular-year ops guide docx from markdown.

Docx 規範（本腳本）：
- 字型：新細明體（PMingLiU）；黑白；禁止彩色／非襯線
- 頁面：A4 直向；左右頁邊距 25mm；頁碼置中「— N —」
- 層級：標題粗體較大＋幼線；內文常規
- 符號：純文字（✓／×）；列表句尾不加句號
"""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "year" / "2627" / "ops-guide.md"
OUT = ROOT / "docs" / "generated" / "2627" / "2627_REGULAR_YEAR_OPS_GUIDE.docx"

FONT_NAME_EA = "新細明體"
BODY_SIZE = 11
H1_SIZE = 14
H2_SIZE = 12
TITLE_SIZE = 18
MARGIN_LR_CM = 2.5  # 25mm
MARGIN_TB_CM = 2.2
INK = RGBColor(0, 0, 0)

EMOJI_MAP = {
    "✅": "✓",
    "❌": "×",
    "⚠": "!",
    "⚠️": "!",
    "→": "→",
}


def normalize_text(text: str) -> str:
    for src, dst in EMOJI_MAP.items():
        text = text.replace(src, dst)
    return text


def strip_list_trailing_period(text: str) -> str:
    """列表／選項句尾不加句號（保留問號、驚嘆號）。"""
    t = text.rstrip()
    while t.endswith("。") or t.endswith("．"):
        t = t[:-1].rstrip()
    return t


def set_run_font(run, size: float, bold: bool = False) -> None:
    run.font.name = FONT_NAME_EA
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = False
    run.font.color.rgb = INK
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT_NAME_EA)
    rFonts.set(qn("w:hAnsi"), FONT_NAME_EA)
    rFonts.set(qn("w:eastAsia"), FONT_NAME_EA)
    # Explicit black (avoid theme colour)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))


def set_style_font(style, size: float, *, bold: bool = False) -> None:
    style.font.name = FONT_NAME_EA
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = False
    style.font.color.rgb = INK
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_NAME_EA)
    rfonts.set(qn("w:hAnsi"), FONT_NAME_EA)
    rfonts.set(qn("w:eastAsia"), FONT_NAME_EA)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)


def add_paragraph_bottom_rule(paragraph, *, size_eighths: int = 6) -> None:
    """標題下方幼線（黑、細）。"""
    pPr = paragraph._p.get_or_add_pPr()
    # Replace existing pBdr if any
    for child in list(pPr):
        if child.tag == qn("w:pBdr"):
            pPr.remove(child)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_margins(cell, *, top=40, bottom=40, left=60, right=60) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_borders(cell, *, bottom_sz: int | None = None) -> None:
    """Black table borders; optional thicker bottom for header hierarchy."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcBorders"):
            tcPr.remove(child)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        sz = bottom_sz if edge == "bottom" and bottom_sz is not None else 4
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "000000")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)


def strip_md_inline(text: str) -> str:
    text = normalize_text(text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


def add_runs_with_bold(paragraph, text: str, *, size: float = BODY_SIZE) -> None:
    text = normalize_text(text)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, bold=True)
        else:
            cleaned = strip_md_inline(part) if "[" in part else part.replace("`", "")
            run = paragraph.add_run(cleaned)
            set_run_font(run, size, bold=False)


def add_para(
    doc: Document,
    text: str,
    *,
    size: float = BODY_SIZE,
    bold: bool = False,
    space_before: float = 0,
    space_after: float = 8,
    align=None,
    bottom_rule: bool = False,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        p.alignment = align
    if bold and "**" not in text:
        run = p.add_run(strip_md_inline(text))
        set_run_font(run, size, bold=True)
    else:
        add_runs_with_bold(p, text, size=size)
    if bottom_rule:
        add_paragraph_bottom_rule(p)


def add_heading1(doc: Document, text: str, *, page_break: bool = True) -> None:
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(strip_md_inline(text), style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.page_break_before = False
    for run in p.runs:
        set_run_font(run, H1_SIZE, bold=True)
    add_paragraph_bottom_rule(p)


def add_heading2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(strip_md_inline(text), style="Heading 2")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    for run in p.runs:
        set_run_font(run, H2_SIZE, bold=True)


def add_toc_title(doc: Document) -> None:
    """目錄標題：粗體＋幼線；非 Heading，唔入 TOC。"""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("目錄")
    set_run_font(run, H1_SIZE, bold=True)
    add_paragraph_bottom_rule(p)


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-1" \h \z \u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "（請於 Word 更新目錄）"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)
    set_run_font(run, BODY_SIZE, bold=False)


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    for child in list(settings):
        if child.tag == qn("w:updateFields"):
            settings.remove(child)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_page_field_run(paragraph) -> None:
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(t)
    r.append(fld_end)
    set_run_font(run, 10, bold=False)


def add_centered_page_number(paragraph) -> None:
    """頁碼格式：— N —"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = paragraph.add_run("— ")
    set_run_font(left, 10, bold=False)
    add_page_field_run(paragraph)
    right = paragraph.add_run(" —")
    set_run_font(right, 10, bold=False)


def add_footer_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        first = section.first_page_footer
        first.is_linked_to_previous = False
        if first.paragraphs:
            first.paragraphs[0].clear()
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        add_centered_page_number(p)


def add_bullet(doc: Document, text: str, *, ordered: bool = False, index: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    prefix = f"{index}. " if ordered else "• "
    run = p.add_run(prefix)
    set_run_font(run, BODY_SIZE, bold=False)
    add_runs_with_bold(p, strip_list_trailing_period(text), size=BODY_SIZE)


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [strip_md_inline(c.strip()) for c in line.split("|")]


def is_table_sep(line: str) -> bool:
    s = line.strip().strip("|").replace(" ", "")
    return bool(s) and all(set(cell) <= {"-", ":"} for cell in s.split("|") if cell)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    norm = [r + [""] * (cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(norm), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    for ri, row in enumerate(norm):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2
            # Header = bold；無底色（黑白印刷）
            run = p.add_run(cell_text)
            set_run_font(run, 10, bold=(ri == 0))
            set_cell_margins(cell)
            set_cell_borders(cell, bottom_sz=12 if ri == 0 else None)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(10)
    run = spacer.add_run("")
    set_run_font(run, BODY_SIZE)


def configure_doc(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(MARGIN_TB_CM)
        section.bottom_margin = Cm(MARGIN_TB_CM)
        section.left_margin = Cm(MARGIN_LR_CM)
        section.right_margin = Cm(MARGIN_LR_CM)
    set_style_font(doc.styles["Normal"], BODY_SIZE)
    for name, size in (("Heading 1", H1_SIZE), ("Heading 2", H2_SIZE)):
        if name in doc.styles:
            style = doc.styles[name]
            set_style_font(style, size, bold=True)
            style.paragraph_format.space_before = Pt(0 if name == "Heading 1" else 14)
            style.paragraph_format.space_after = Pt(12 if name == "Heading 1" else 8)
            style.paragraph_format.page_break_before = False


def add_cover(doc: Document, meta: dict[str, str]) -> None:
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_before = Pt(48)
    brand.paragraph_format.space_after = Pt(8)
    run = brand.add_run("明學教育")
    set_run_font(run, 14, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("2627 常規學年營運指引")
    set_run_font(run, TITLE_SIZE, bold=True)
    add_paragraph_bottom_rule(title, size_eighths=8)

    cover_rows = [
        ["項目", "內容"],
        ["適用學年", meta.get("適用學年", "")],
        ["文件版本", meta.get("文件版本", "")],
        ["更新日期", meta.get("更新日期", "")],
        ["對象", meta.get("對象", "")],
        ["發佈", meta.get("發佈", "")],
    ]
    # spacing before meta table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(12)
    spacer.paragraph_format.space_after = Pt(0)
    add_table(doc, cover_rows)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(16)
    run = note.add_run("全公司員工守則｜列印／WhatsApp 發佈用")
    set_run_font(run, 10, bold=False)


def convert(md_text: str) -> Document:
    doc = Document()
    configure_doc(doc)

    lines = md_text.splitlines()
    i = 0
    if lines and lines[0].startswith("# "):
        i = 1

    meta: dict[str, str] = {}
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines) and lines[i].strip().startswith("|"):
        table_lines: list[str] = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i])
            i += 1
        rows = [parse_table_row(l) for l in table_lines if not is_table_sep(l)]
        for row in rows[1:]:
            if len(row) >= 2:
                meta[row[0]] = row[1]

    add_cover(doc, meta)

    while i < len(lines) and (not lines[i].strip() or lines[i].strip() == "---"):
        i += 1

    skip_md_toc_body = False
    ordered_index = 0

    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            i += 1
            ordered_index = 0
            continue

        if stripped == "---":
            skip_md_toc_body = False
            i += 1
            ordered_index = 0
            continue

        if stripped.startswith("## "):
            title = stripped[3:]
            if title == "目錄":
                add_toc_title(doc)
                add_toc_field(doc)
                skip_md_toc_body = True
                i += 1
                ordered_index = 0
                continue
            skip_md_toc_body = False
            add_heading1(doc, title, page_break=True)
            i += 1
            ordered_index = 0
            continue

        if skip_md_toc_body:
            i += 1
            continue

        if stripped.startswith("### "):
            sub = stripped[4:]
            if sub.strip() == "重點":
                add_para(doc, "重點", bold=True, size=H2_SIZE, space_before=14, space_after=8)
            else:
                add_heading2(doc, sub)
            i += 1
            ordered_index = 0
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [parse_table_row(l) for l in table_lines if not is_table_sep(l)]
            add_table(doc, rows)
            ordered_index = 0
            continue

        m_ol = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_ol:
            ordered_index = int(m_ol.group(1))
            add_bullet(doc, m_ol.group(2), ordered=True, index=ordered_index)
            i += 1
            continue

        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:], ordered=False)
            i += 1
            ordered_index = 0
            continue

        add_para(doc, stripped)
        i += 1
        ordered_index = 0

    add_footer_page_numbers(doc)
    # 唔設 updateFields=true：Word 開啟時會彈對話框卡住自動化
    return doc


def close_word_docs() -> None:
    subprocess.run(
        [
            "osascript",
            "-e",
            'tell application "Microsoft Word"\n'
            "  repeat with d in (get documents)\n"
            "    try\n"
            "      close d saving no\n"
            "    end try\n"
            "  end repeat\n"
            "end tell",
        ],
        check=False,
        capture_output=True,
        text=True,
    )


def update_word_toc(path: Path) -> None:
    """Refresh Word built-in TOC (and fields) via Microsoft Word on macOS."""
    script = f'''
tell application "Microsoft Word"
  activate
  set docPath to POSIX file "{path}" as alias
  open docPath
  delay 1
  set theDoc to active document
  try
    if (count of tables of contents of theDoc) > 0 then
      update table of contents 1 of theDoc
    end if
  end try
  set n to count of fields of theDoc
  repeat with i from 1 to n
    try
      update field i of theDoc
    end try
  end repeat
  delay 0.5
  save theDoc
end tell
'''
    subprocess.run(["osascript", "-e", script], check=False, capture_output=True, text=True)


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Generate 2627 ops guide docx")
    parser.add_argument(
        "--no-word",
        action="store_true",
        help="Skip Microsoft Word TOC/field refresh",
    )
    args = parser.parse_args()

    if not args.no_word:
        close_word_docs()
    md = SRC.read_text(encoding="utf-8")
    doc = convert(md)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("wrote", OUT)
    if args.no_word:
        print("skipped Word TOC update (--no-word)")
        return
    try:
        update_word_toc(OUT)
        print("updated Word TOC / fields")
    except Exception as exc:  # noqa: BLE001
        print("Word TOC update skipped:", exc)
        print("Open in Word → 右鍵目錄 → 更新功能變數")


if __name__ == "__main__":
    main()
