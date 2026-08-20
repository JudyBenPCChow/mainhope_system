#!/usr/bin/env python3
"""Generate 2627 empty-room list + 日視圖1 from the current timetable scheme.

Default: markdown. Pass --word to also write docx (no PDF).
日視圖1 = 朝 9 至晚 6（09:00-17:45）；晚間 17:45 起不列。
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPTS_DIR = Path(__file__).resolve().parent
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import generate_2627_timetable_doc as tt

STEM_EMPTY = f"2627_timetable_empty_rooms_v{tt.VERSION}"
EMPTY_EVEN_HEADER = f"明學教育 2627學年空房時間（ver {tt.VERSION}）"
DAYTIME_LAST_SLOT = 6  # 16:30-17:45；與系統日視圖「朝 9–晚 6」一致
NOTE_SCHOOL = "返學時間"
NOTE_NO_CLASS = "不排課"
EMPTY_CELL_LABELS = {"—", NOTE_NO_CLASS}


def is_empty_cell(cell: str) -> bool:
    return cell in EMPTY_CELL_LABELS


def is_school_hours_row(row: list[str]) -> bool:
    return all(cell == NOTE_SCHOOL for cell in row[1:-1])


def cell_note(cell: str) -> str:
    if cell == NOTE_NO_CLASS:
        return NOTE_NO_CLASS
    return ""


def empty_rooms_in_row(row: list[str]) -> tuple[list[str], str, int]:
    rooms: list[str] = []
    notes: set[str] = set()
    for room, cell in zip(tt.ROOMS, row[1:-1]):
        if cell == NOTE_SCHOOL:
            continue
        if is_empty_cell(cell):
            rooms.append(room)
            note = cell_note(cell)
            if note:
                notes.add(note)
    note = notes.pop() if len(notes) == 1 else "、".join(sorted(notes))
    return rooms, note, len(rooms)


def daytime_grid(day_idx: int) -> list[list[str]]:
    raw = tt.day_grid(day_idx)
    grid = [raw[0]]
    for row in raw[1 : DAYTIME_LAST_SLOT + 2]:
        _rooms, _note, counted = empty_rooms_in_row(row)
        grid.append([*row[:-1], str(counted)])
    return grid


def slot_span_label(start_slot: int, end_slot: int) -> str:
    start = tt.SLOT_TIMES[start_slot].split("-")[0]
    end = tt.SLOT_TIMES[end_slot].split("-")[1]
    return f"{start}-{end}"


def day_empty_ranges(day_idx: int) -> list[tuple[str, list[str], str, int]]:
    grid = daytime_grid(day_idx)
    ranges: list[tuple[str, list[str], str, int]] = []
    pending: tuple[int, int, tuple[str, ...], str, int] | None = None
    for slot_idx, row in enumerate(grid[1:]):
        if is_school_hours_row(row):
            continue
        rooms, note, counted = empty_rooms_in_row(row)
        if counted == 0:
            continue
        key = (tuple(rooms), note, counted)
        if pending and (pending[2], pending[3], pending[4]) == key:
            pending = (pending[0], slot_idx, pending[2], pending[3], pending[4])
        else:
            if pending:
                ranges.append(
                    (slot_span_label(pending[0], pending[1]), list(pending[2]), pending[3], pending[4])
                )
            pending = (slot_idx, slot_idx, key[0], key[1], key[2])
    if pending:
        ranges.append((slot_span_label(pending[0], pending[1]), list(pending[2]), pending[3], pending[4]))
    return ranges


def rooms_text(rooms: list[str]) -> str:
    return "、".join(rooms) if rooms else "無"


def list_detail_rows() -> list[list[str]]:
    rows = [["星期", "時段", "空房", "備註", "空房數"]]
    for day_idx, day_name in enumerate(tt.DAYS):
        grid = daytime_grid(day_idx)
        for row in grid[1:]:
            if is_school_hours_row(row):
                continue
            rooms, note, counted = empty_rooms_in_row(row)
            if counted == 0:
                continue
            rows.append([day_name, row[0], rooms_text(rooms), note, str(counted)])
    return rows


def cover_notes() -> list[str]:
    return [
        f"本文件對應方案紀錄 ver. {tt.VERSION}，尚未寫入正式班別與排程",
        "列表與日視圖1均只列朝 9 至晚 6（09:00-17:45）；晚間 17:45-20:15 見周時間表",
        "空房＝該課室該時段無常規班、無預留時段、亦非功課輔導班專用",
        "平日返學時間（15:15-16:30 或之前）不列入空房列表，亦不計空房數",
        "週末 09:00-10:15 標不排課，該列仍計空房",
        "平日 16:30 起 17D 為功課輔導班專用，不計空房",
        "不排 20:15-21:30，本檔亦不列該格",
    ]


def range_bullet(span: str, rooms: list[str], note: str, counted: int) -> str:
    suffix = f"（{note}；空房 {counted}）" if note else f"（空房 {counted}）"
    return f"{span}：{rooms_text(rooms)}{suffix}"


def build_empty_rooms_md() -> str:
    parts: list[str] = [
        "# 明學教育 2627 學年常規專科班空房時間",
        "",
        f"列表及日視圖1｜對應方案紀錄 **ver. {tt.VERSION}**｜未定稿入庫",
        "",
        "> 本檔由 `python3 scripts/generate_2627_empty_rooms_doc.py` 生成；出 Word 時加 `--word`。",
        "",
        "## 封面說明",
        tt.md_bullets(cover_notes()),
        "",
        "## 1. 列表",
        "",
        "以下按日合併連續時段。返學時間不列入。五室皆空時仍逐室列出，以便對照日視圖1。",
    ]
    for day_idx, day_name in enumerate(tt.DAYS):
        parts.extend(["", f"### 1.{day_idx + 1} {day_name}", ""])
        bullets = [
            range_bullet(span, rooms, note, counted)
            for span, rooms, note, counted in day_empty_ranges(day_idx)
        ]
        parts.append(tt.md_bullets(bullets) if bullets else "- 本日日視圖1無空房")
    parts.extend(
        [
            "",
            "### 1.8 明細表",
            "",
            "以下不合併時段，一列一格。返學時間不列入。",
            "",
            tt.md_table(list_detail_rows()),
            "",
            "## 2. 日視圖1",
            "",
            tt.md_bullets(
                [
                    "日視圖1＝朝 9 至晚 6（09:00-17:45），共 7 格",
                    "格內「—」為空房；返學時間不計空房（空房數 0）；不排課仍計空房；功課輔導班不計空房",
                    "最右欄為該時段空房數",
                ]
            ),
        ]
    )
    for day_idx, day_name in enumerate(tt.DAYS):
        parts.extend(
            [
                "",
                f"### 2.{day_idx + 1} {day_name}",
                "",
                tt.md_table(daytime_grid(day_idx)),
            ]
        )
    parts.extend(["", "— 完 —", ""])
    return "\n".join(parts)


def write_markdown() -> Path:
    tt.OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = tt.OUT_DIR / f"{STEM_EMPTY}.md"
    path.write_text(build_empty_rooms_md(), encoding="utf-8")
    return path


def add_daytime_table(doc: Document, day_idx: int, widths: list[float]) -> None:
    grid = daytime_grid(day_idx)
    table = doc.add_table(rows=len(grid), cols=len(grid[0]))
    table.style = "Table Grid"
    last_col = len(grid[0]) - 1
    for i, row_vals in enumerate(grid):
        tt.set_row_height_exact(table.rows[i], tt.WEEKLY_DATA_ROW_CM if i else tt.WEEKLY_HEADER_ROW_CM)
        for j, val in enumerate(row_vals):
            muted = val in tt.MUTED_LABELS
            empty = val == "—"
            tt.write_cell(
                table.rows[i].cells[j],
                val,
                bold=(i == 0 or j == 0 or j == last_col),
                size=8 if i else 10,
                color=tt.COLOR_MUTED if muted else None,
            )
            tt.set_cell_valign(table.rows[i].cells[j])
            if i == 0:
                tt.shade_cell(table.rows[i].cells[j])
            elif muted:
                tt.shade_cell(table.rows[i].cells[j], tt.FILL_SCHOOL)
            elif empty:
                tt.shade_cell(table.rows[i].cells[j], "F3FFF3")
            if j == 0 or j == last_col:
                for para in table.rows[i].cells[j].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tt.set_cell_nowrap(table.rows[i].cells[j])
    tt.set_table_widths(table, widths)
    tt.prevent_row_split(table)


def build_empty_rooms_docx(path: Path) -> None:
    doc = Document()
    tt.configure_docx_styles(doc)
    tt.enable_even_odd_headers(doc)
    tt.apply_section_chrome(doc.sections[0], EMPTY_EVEN_HEADER, first_empty=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("明學教育")
    tt.set_run_font(run, 12, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("2627 學年常規專科班空房時間")
    tt.set_run_font(run, 12, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"列表及日視圖1｜對應方案紀錄 ver. {tt.VERSION}｜未定稿入庫")
    tt.set_run_font(run, 12)

    tt.add_bullets(doc, cover_notes())
    tt.add_page_break(doc)
    tt.add_para(doc, "目錄", size=tt.H1_SIZE, bold=True, space_after=8)
    tt.add_toc_field(doc)

    tt.add_chapter(doc, "1. 列表")
    tt.add_para(doc, "以下按日合併連續時段。返學時間不列入。五室皆空時仍逐室列出，以便對照日視圖1。", size=12)
    for day_idx, day_name in enumerate(tt.DAYS):
        tt.add_section_title(doc, f"1.{day_idx + 1} {day_name}", keep_with_next=True)
        bullets = [
            range_bullet(span, rooms, note, counted)
            for span, rooms, note, counted in day_empty_ranges(day_idx)
        ]
        tt.add_bullets(doc, bullets if bullets else ["本日日視圖1無空房"])
    tt.add_section_title(doc, "1.8 明細表", keep_with_next=True)
    tt.add_para(doc, "以下不合併時段，一列一格。返學時間不列入。", size=12)
    tt.add_docx_data_table(doc, list_detail_rows(), size=10)

    tt.add_chapter(doc, "2. 日視圖1")
    tt.add_bullets(
        doc,
        [
            "日視圖1＝朝 9 至晚 6（09:00-17:45），共 7 格",
            "格內「—」為空房（淺綠）；返學時間不計空房（空房數 0）；不排課仍計空房；功課輔導班不計空房",
            "最右欄為該時段空房數；每日一頁、橫向",
        ],
    )

    land = doc.add_section(WD_SECTION.NEW_PAGE)
    tt.apply_normal_page(land, landscape=True)
    tt.apply_section_chrome(land, EMPTY_EVEN_HEADER, first_empty=False)

    usable = tt.A4_H_CM - tt.MARGIN_CM * 2
    time_w, empty_w = 2.15, 1.2
    room_w = (usable - time_w - empty_w) / 5
    widths = [time_w] + [room_w] * 5 + [empty_w]
    for day_idx, day_name in enumerate(tt.DAYS):
        tt.add_heading(
            doc,
            f"2.{day_idx + 1} {day_name}",
            2,
            page_break_before=(day_idx > 0),
            keep_with_next=True,
        )
        add_daytime_table(doc, day_idx, widths)

    tt.add_para(doc, "— 完 —", size=12)
    doc.save(path)


def validate_empty_rooms() -> None:
    tt.validate()
    for day_idx in range(len(tt.DAYS)):
        grid = daytime_grid(day_idx)
        assert len(grid) == DAYTIME_LAST_SLOT + 2
        assert grid[0][-1] == "空房"
        assert grid[-1][0] == tt.SLOT_TIMES[DAYTIME_LAST_SLOT]
        for row in grid[1:]:
            if is_school_hours_row(row):
                assert row[-1] == "0"


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate 2627 empty-room md (default) or Word")
    parser.add_argument("--word", action="store_true", help="Also write docx")
    args = parser.parse_args()
    validate_empty_rooms()
    md_path = write_markdown()
    print("wrote", md_path)
    if not args.word:
        print("skip docx (pass --word to export Word)")
        return
    docx_path = tt.OUT_DIR / f"{STEM_EMPTY}.docx"
    build_empty_rooms_docx(docx_path)
    print("wrote", docx_path)


if __name__ == "__main__":
    main()
