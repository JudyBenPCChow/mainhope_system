#!/usr/bin/env python3
"""Generate an editable Word copy of the complete MainHope terminology table.

Source: docs/meta/TERMINOLOGY.md
Output: docs/generated/frontdesk/MAINHOPE_TERMINOLOGY.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "meta" / "TERMINOLOGY.md"
OUTPUT = ROOT / "docs" / "generated" / "frontdesk" / "MAINHOPE_TERMINOLOGY.docx"

FONT_NAME = "新細明體"
INK = RGBColor(0, 0, 0)


def clean_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)


def clean_inline(text: str) -> str:
    return clean_links(text).replace("**", "").replace("`", "").strip()


def set_run_font(run, size: float, *, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = INK
    properties = run._element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    language = OxmlElement("w:lang")
    language.set(qn("w:eastAsia"), "zh-HK")
    properties.append(language)


def set_style_font(style, size: float, *, bold: bool = False) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = INK
    properties = style.element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)


def add_inline_runs(paragraph, text: str, *, size: float = 10.5) -> None:
    text = clean_links(text).replace("`", "")
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if is_bold else part)
        set_run_font(run, size, bold=is_bold)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run("明學教育｜完整用語表　— ")
    set_run_font(prefix, 9)
    field = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field._r.extend([begin, instruction, separate, text, end])
    set_run_font(field, 9)
    suffix = paragraph.add_run(" —")
    set_run_font(suffix, 9)


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    properties = cell._tc.get_or_add_tcPr()
    width = properties.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(int(width_cm * 567)))
    width.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def parse_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(set(cell) <= {"-", ":", " "} for cell in cells)


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for row_index, source_row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        if row_index == 0:
            repeat_header(table.rows[-1])
        for column_index, value in enumerate(source_row):
            cell = cells[column_index]
            set_cell_width(cell, 4.3 if column_index == 0 else 12.9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline_runs(paragraph, value, size=9.5)
            if row_index == 0:
                for run in paragraph.runs:
                    run.font.bold = True
                shade_cell(cell, "D9D9D9")
            elif row_index % 2 == 0:
                shade_cell(cell, "F2F2F2")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.footer_distance = Cm(0.8)

    set_style_font(document.styles["Normal"], 10.5)
    set_style_font(document.styles["Title"], 20, bold=True)
    set_style_font(document.styles["Heading 1"], 14, bold=True)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def build_document() -> None:
    raw_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    update_date = next(
        line.removeprefix("更新日期：").strip()
        for line in raw_lines
        if line.startswith("更新日期：")
    )
    lines = [line for line in raw_lines if not line.startswith("更新日期：")]

    document = Document()
    configure_document(document)
    document.core_properties.title = "明學教育完整用語表"
    document.core_properties.author = "明學教育"
    document.core_properties.subject = "公司用詞及其含義"

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    add_inline_runs(title, clean_inline(lines[0].removeprefix("# ")), size=20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run(f"全體員工查閱｜更新日期：{update_date}")
    set_run_font(subtitle_run, 10)

    index = 1
    while index < len(lines):
        line = lines[index].strip()
        if not line or line == "---":
            index += 1
            continue
        if line.startswith("## "):
            heading_text = clean_inline(line[3:])
            paragraph = document.add_paragraph(style="Heading 1")
            if line.startswith(("## 4.", "## 7.", "## 10.")):
                paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(heading_text)
            set_run_font(run, 14, bold=True)
            index += 1
            continue
        if line.startswith("|"):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                row = parse_table_row(lines[index])
                if not is_separator_row(row):
                    rows.append(row)
                index += 1
            add_table(document, rows)
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline_runs(paragraph, line[2:])
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        add_inline_runs(paragraph, line)
        index += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    build_document()
