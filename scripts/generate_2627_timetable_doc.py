#!/usr/bin/env python3
"""Generate 2627 timetable records.

Default: markdown review drafts for the scheme, teacher appendix, and standalone
weekly timetable. Pass --word to write matching docx files (Word built-in TOC /
header / footer / Normal margins) and export PDFs via Microsoft Word (not reportlab).
"""

from __future__ import annotations

import argparse
import csv
import subprocess
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TIMETABLE_DIR = ROOT / "docs" / "year" / "2627" / "timetable"
VERSION = "4.0"
PREV_VERSION = "3.10"
STEM = f"2627_timetable_scheme_v{VERSION}"
STEM_TEACHERS = f"2627_timetable_teachers_week_v{VERSION}"
STEM_WEEKLY = f"2627_timetable_weekly_v{VERSION}"
STEM_CODES = f"2627_timetable_class_codes_v{VERSION}"
OUT_DIR = TIMETABLE_DIR / "versions" / f"v{VERSION}"
# 改 CLASSES／原則／已確認老師鎖時把 VERSION 改成下一 4.x（4.0→4.1），PREV_VERSION＝舊版；舊檔保留。
FONT_NAME_EA = "新細明體"
MARGIN_CM = 2.54  # Word 預設「普通」
HEADER_DISTANCE_CM = 1.25
FOOTER_DISTANCE_CM = 1.25
A4_W_CM = 21.0
A4_H_CM = 29.7
H1_SIZE = 14
H2_SIZE = 12
H3_SIZE = 12

SLOTS = [
    "09:00–10:15",
    "10:15–11:30",
    "11:30–12:45",
    "12:45–14:00",
    "14:00–15:15",
    "15:15–16:30",
    "16:30–17:45",
    "17:45–19:00",
    "19:00–20:15",
]
ROOMS = ["17D", "17E", "矩尺座", "英仙座", "山案座"]
DAYS = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]

GRADES = ["S1", "S2", "S3", "S4", "S5", "S6"]
GRADE_FULL = {
    "S1": "中一級",
    "S2": "中二級",
    "S3": "中三級",
    "S4": "中四級",
    "S5": "中五級",
    "S6": "中六級",
}
SUBJECT_FULL = {
    "中文": "中文科",
    "英文": "英文科",
    "數學": "數學科",
    "科學": "科學科",
    "物理": "物理科",
    "化學": "化學科",
    "生物": "生物科",
    "企會財": "企會財科",
}
# Display times in timetable cells (ASCII hyphen, per record format).
SLOT_TIMES = [
    "09:00-10:15",
    "10:15-11:30",
    "11:30-12:45",
    "12:45-14:00",
    "14:00-15:15",
    "15:15-16:30",
    "16:30-17:45",
    "17:45-19:00",
    "19:00-20:15",
]


def class_section_letter(code: str) -> str:
    return code[-1]


def class_title(grade: str, subject: str, code: str, teacher: str | None = None) -> str:
    """e.g. 中六級中文科（A)"""
    if teacher and (teacher, grade) in JACKSON_MIXED_ENG_BASE and subject == "英文":
        return f"{MIXED_SENIOR_ENG_TITLE}（{class_section_letter(code)})"
    return f"{GRADE_FULL[grade]}{SUBJECT_FULL[subject]}（{class_section_letter(code)})"


def class_cell_text(grade: str, subject: str, teacher: str, code: str, slot_idx: int) -> str:
    return f"{class_title(grade, subject, code, teacher)}\n{teacher}\n{SLOT_TIMES[slot_idx]}"


SUBJECT_TO_COURSE_ABBR = {
    "中文": "CHI",
    "英文": "ENG",
    "數學": "MATH",
    "科學": "SCI",
    "物理": "PHY",
    "化學": "CHEM",
    "生物": "BIO",
}
# 課程名稱跟系統 courses.course_name 慣例：中X級常規Y班（中文課省略「級」）。
COURSE_TEMPLATE_NAMES = {
    "BIOS4001": "中四級常規生物班",
    "BIOS5001": "中五級常規生物班",
    "BIOS6001": "中六級常規生物班",
    "CHIS1001": "中一常規中文班",
    "CHIS2001": "中二常規中文班",
    "CHIS3001": "中三常規中文班",
    "CHIS4001": "中四常規中文班",
    "CHIS5001": "中五常規中文班",
    "CHIS6001": "中六常規中文班",
    "ENGS1001": "中一級常規英文班",
    "ENGS2001": "中二級常規英文班",
    "ENGS3001": "中三級常規英文班",
    "ENGS4001": "中四級常規英文班",
    "ENGS4004": "高中常規英文班（中四至中六）",
    "ENGS5001": "中五級常規英文班",
    "ENGS5004": "高中常規英文班（中四至中六）",
    "ENGS6001": "中六級常規英文班",
    "MATHS1001": "中一級常規數學班",
    "MATHS2001": "中二級常規數學班",
    "MATHS3001": "中三級常規數學班",
    "MATHS4001": "中四級常規數學班",
    "MATHS5001": "中五級常規數學班",
    "MATHS6001": "中六級常規數學班",
    "PHYS4001": "中四級常規物理班",
    "PHYS5001": "中五級常規物理班",
    "PHYS6001": "中六級常規物理班",
    "SCIS1001": "中一級常規科學班",
    "SCIS2001": "中二級常規科學班",
    "SCIS3001": "中三級常規科學班",
    "CHEMS4001": "中四級常規化學班",
    "CHEMS5001": "中五級常規化學班",
    "CHEMS6001": "中六級常規化學班",
}


def course_code_base(subject: str, grade: str) -> str:
    return f"{SUBJECT_TO_COURSE_ABBR[subject]}{grade}001"


# Jackson Lau 2627 兩班英文：高中混級模板（seq 004），非該級 001。
JACKSON_MIXED_ENG_BASE = {
    ("Jackson Lau", "S4"): "ENGS4004",
    ("Jackson Lau", "S5"): "ENGS5004",
}
MIXED_SENIOR_ENG_TITLE = "高中英文科（中四至中六）"


def course_code_full(subject: str, grade: str, section: str, teacher: str | None = None) -> str:
    base = JACKSON_MIXED_ENG_BASE.get((teacher or "", grade)) or course_code_base(subject, grade)
    return f"2627-{base}-{section}"


def assign_chinese_class_letters(
    classes: list[tuple[int, int, str, str, str, str, str]],
) -> list[tuple[int, int, str, str, str, str, str]]:
    """各級中文班號按星期→時段重編為連續 A、B、C…。"""
    by_grade: dict[str, list[int]] = defaultdict(list)
    for i, row in enumerate(classes):
        if row[3] == "中文":
            by_grade[row[4]].append(i)
    out = list(classes)
    for grade, idxs in by_grade.items():
        idxs.sort(key=lambda i: (out[i][0], out[i][1], out[i][2]))
        for n, i in enumerate(idxs):
            day_idx, slot_idx, room, subject, g, teacher, _code = out[i]
            letter = chr(ord("A") + n)
            out[i] = (day_idx, slot_idx, room, subject, g, teacher, f"{g}中{letter}")
    return out


# (day_idx, slot_idx, room, subject, grade, teacher, code)
# ver. 3.2：以 3.1 為底，Henry 只排星期六連續三堂中四／中五／中六生物。
# 中文班號於表末由 assign_chinese_class_letters 覆寫；此處碼僅佔位。
CLASSES: list[tuple[int, int, str, str, str, str, str]] = [
    # Monday — Mark 矩尺連三；Katie 17E 兩堂（本版不排 19:00）
    (0, 6, "矩尺座", "數學", "S1", "Mark Yu", "S1數A"),
    (0, 6, "17E", "中文", "S2", "Katie", "S2中"),
    (0, 7, "矩尺座", "數學", "S5", "Mark Yu", "S5數B"),
    (0, 7, "17E", "中文", "S3", "Katie", "S3中"),
    (0, 8, "矩尺座", "數學", "S4", "Mark Yu", "S4數B"),
    # Tuesday — Mark 矩尺連三；Katie 17E 兩堂
    (1, 6, "矩尺座", "數學", "S2", "Mark Yu", "S2數A"),
    (1, 6, "17E", "中文", "S3", "Katie", "S3中"),
    (1, 7, "矩尺座", "數學", "S6", "Mark Yu", "S6數B"),
    (1, 7, "17E", "中文", "S1", "Katie", "S1中"),
    (1, 8, "矩尺座", "數學", "S4", "Mark Yu", "S4數A"),
    # Wednesday — 無 Mark；Katie 留矩尺兩堂；Jackson 山案（不排 17D／17E）
    (2, 6, "矩尺座", "中文", "S1", "Katie", "S1中"),
    (2, 7, "矩尺座", "中文", "S2", "Katie", "S2中"),
    (2, 7, "山案座", "英文", "S5", "Jackson Lau", "S5英B"),
    # Thursday — Mark 矩尺高中兩班；Katie 17E 兩堂
    (3, 6, "17E", "中文", "S3", "Katie", "S3中"),
    (3, 7, "矩尺座", "數學", "S5", "Mark Yu", "S5數A"),
    (3, 7, "17E", "中文", "S1", "Katie", "S1中"),
    (3, 8, "矩尺座", "數學", "S6", "Mark Yu", "S6數A"),
    # Friday — Judy 中五生物矩尺；Christine 本版不排星期五
    (4, 7, "矩尺座", "生物", "S5", "Judy Chu", "S5生A"),
    # Saturday — Mark 矩尺上午兩堂、12:45 午膳（不標）、14:00／15:15／17:45；Jackson 12:45 英仙（矩尺讓 Christine）；Leo 山案；Christine 12:45 矩尺中四、14:00 山案中五、16:30 矩尺中六（本版不避撞科）
    (5, 1, "矩尺座", "數學", "S3", "Mark Yu", "S3數A"),
    (5, 1, "山案座", "數學", "S1", "Leo Chan", "S1數B"),
    (5, 1, "17E", "英文", "S2", "Cheryl Ng", "S2英B"),
    (5, 1, "17D", "化學", "S5", "Phoebe Tam", "S5化A"),
    (5, 2, "矩尺座", "數學", "S5", "Mark Yu", "S5數D"),
    (5, 2, "山案座", "物理", "S4", "Leo Chan", "S4物A"),
    (5, 2, "17E", "英文", "S1", "Cheryl Ng", "S1英B"),
    (5, 2, "17D", "科學", "S2", "Phoebe Tam", "S2科A"),
    (5, 3, "英仙座", "英文", "S4", "Jackson Lau", "S4英B"),
    (5, 3, "山案座", "數學", "S2", "Leo Chan", "S2數C"),
    (5, 3, "17D", "化學", "S6", "Phoebe Tam", "S6化A"),
    (5, 3, "矩尺座", "中文", "S4", "Christine Fan", "S4中"),
    (5, 4, "矩尺座", "數學", "S6", "Mark Yu", "S6數D"),
    (5, 4, "17E", "數學", "S2", "Liam Lai", "S2數D"),
    (5, 4, "英仙座", "生物", "S4", "Henry Wong", "S4生A"),
    (5, 4, "山案座", "中文", "S5", "Christine Fan", "S5中"),
    (5, 5, "矩尺座", "數學", "S1", "Mark Yu", "S1數C"),
    (5, 5, "17E", "數學", "S3", "Liam Lai", "S3數C"),
    (5, 5, "山案座", "物理", "S6", "Leo Chan", "S6物A"),
    (5, 5, "英仙座", "生物", "S5", "Henry Wong", "S5生B"),
    (5, 5, "17D", "化學", "S4", "Phoebe Tam", "S4化A"),
    (5, 6, "山案座", "物理", "S5", "Leo Chan", "S5物A"),
    (5, 6, "英仙座", "生物", "S6", "Henry Wong", "S6生C"),
    (5, 6, "17E", "中文", "S3", "Billy Shek", "S3中"),
    (5, 6, "矩尺座", "中文", "S6", "Christine Fan", "S6中"),
    (5, 7, "17E", "中文", "S2", "Billy Shek", "S2中"),
    (5, 7, "矩尺座", "數學", "S4", "Mark Yu", "S4數C"),
    # Sunday — Katie 17E 五堂；Christine 矩尺三堂；Cyndi 英仙；Emma 17D
    (6, 1, "17E", "中文", "S1", "Katie", "S1中"),
    (6, 1, "英仙座", "英文", "S6", "Cyndi Ng", "S6英A"),
    (6, 1, "山案座", "數學", "S2", "Liam Lai", "S2數B"),
    (6, 1, "矩尺座", "科學", "S3", "Phoebe Tam", "S3科A"),
    (6, 2, "17E", "中文", "S2", "Katie", "S2中"),
    (6, 2, "英仙座", "英文", "S5", "Cyndi Ng", "S5英A"),
    (6, 2, "山案座", "數學", "S3", "Liam Lai", "S3數B"),
    (6, 2, "17D", "生物", "S6", "Judy Chu", "S6生A"),
    (6, 2, "矩尺座", "科學", "S1", "Phoebe Tam", "S1科A"),
    (6, 3, "矩尺座", "中文", "S5", "Christine Fan", "S5中"),
    (6, 3, "17D", "英文", "S2", "Emma Cai", "S2英A"),
    (6, 3, "17E", "生物", "S6", "Judy Chu", "S6生B"),
    (6, 4, "17E", "中文", "S3", "Katie", "S3中"),
    (6, 4, "英仙座", "英文", "S4", "Cyndi Ng", "S4英A"),
    (6, 4, "17D", "英文", "S1", "Emma Cai", "S1英A"),
    (6, 5, "17E", "中文", "S2", "Katie", "S2中"),
    (6, 5, "矩尺座", "中文", "S4", "Christine Fan", "S4中"),
    (6, 6, "17E", "中文", "S1", "Katie", "S1中"),
    (6, 6, "矩尺座", "中文", "S6", "Christine Fan", "S6中"),
    (6, 6, "17D", "英文", "S3", "Emma Cai", "S3英A"),
    (6, 7, "17D", "英文", "S3", "Emma Cai", "S3英B"),
]
CLASSES = assign_chinese_class_letters(CLASSES)

# (day_idx, slot_idx, room, teacher, title)
# Reserved non-group slots shown on the timetable.
RESERVED: list[tuple[int, int, str, str, str]] = [
    (6, 5, "英仙座", "Cyndi Ng", "一對一高中英文科（預留）"),
]

# 已確認老師的現行班別時間；除非用戶明確指示，後續方案不可改動。
# 值為 (day_idx, slot_idx, class_code)，不包括預留時段。
CONFIRMED_TEACHER_CLASS_TIMES: dict[str, set[tuple[int, int, str]]] = {
    "Cyndi Ng": {
        (6, 1, "S6英A"),
        (6, 2, "S5英A"),
        (6, 4, "S4英A"),
    },
    "Emma Cai": {
        (6, 3, "S2英A"),
        (6, 4, "S1英A"),
        (6, 6, "S3英A"),
        (6, 7, "S3英B"),
    },
    "Liam Lai": {
        (5, 4, "S2數D"),
        (5, 5, "S3數C"),
        (6, 1, "S2數B"),
        (6, 2, "S3數B"),
    },
    "Leo Chan": {
        (5, 1, "S1數B"),
        (5, 2, "S4物A"),
        (5, 3, "S2數C"),
        (5, 5, "S6物A"),
        (5, 6, "S5物A"),
    },
}

CLASS_CODES_FIELDS = [
    "班別顯示碼（course_code_full）",
    "課程名稱（course_name）",
    "班號（section_code）",
    "班別",
    "任教老師",
    "星期",
    "時段",
    "課室",
]


def class_codes_rows() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    ordered = sorted(CLASSES, key=lambda c: (c[0], c[1], c[2]))
    for day_idx, slot_idx, room, subject, grade, teacher, code in ordered:
        section = class_section_letter(code)
        mixed_base = JACKSON_MIXED_ENG_BASE.get((teacher, grade))
        base = mixed_base or course_code_base(subject, grade)
        title = (
            MIXED_SENIOR_ENG_TITLE
            if mixed_base
            else f"{GRADE_FULL[grade]}{SUBJECT_FULL[subject]}"
        )
        rows.append(
            {
                "班別顯示碼（course_code_full）": course_code_full(subject, grade, section, teacher),
                "課程名稱（course_name）": COURSE_TEMPLATE_NAMES[base],
                "班號（section_code）": section,
                "班別": f"{title}（{section}）",
                "任教老師": teacher,
                "星期": DAYS[day_idx],
                "時段": SLOT_TIMES[slot_idx],
                "課室": room,
            }
        )
    return rows


def write_class_codes_csv() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / f"{STEM_CODES}.csv"
    rows = class_codes_rows()
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=CLASS_CODES_FIELDS)
        writer.writeheader()
        writer.writerows(rows)
    data = path.read_bytes()
    assert data.startswith(b"\xef\xbb\xbf"), path
    with path.open(encoding="utf-8-sig", newline="") as f:
        check = list(csv.DictReader(f))
    assert len(check) == len(CLASSES)
    assert len({row["班別顯示碼（course_code_full）"] for row in check}) == len(CLASSES)
    return path


MINUTES_PER_CLASS = 75

STAFF = [
    (
        "Mark Yu",
        "數學科",
        13,
        "星期一、星期二、星期四、星期六",
        "兼職。必須星期六出勤；不排星期三、星期五、星期日。平日每日最多三班；週末每日最多五班。本版十三班：星期一、二連排三堂；星期四高中兩班；星期六五班（10:15、11:30；12:45 午膳不標示；14:00、15:15、17:45；16:30 矩尺讓 Christine）。出勤日優先矩尺座。",
    ),
    (
        "Katie",
        "中文科",
        13,
        "星期一至星期四、星期日",
        "全職。放假星期五、星期六。平日 14:00 至最後一節；週末 09:00-18:00，中間一節食飯休息。本版十三班：星期一至四每日兩堂（16:30、17:45）；星期日五班（10:15 至 16:30，12:45 食飯休息）。老師附件：非上堂時間標空堂；放假日灰底。",
    ),
    (
        "Christine Fan",
        "中文科（中四級至中六級）",
        6,
        "星期六、星期日",
        "兼職。必須星期日出勤。本版只排星期六、星期日共六班。星期六：12:45 中四（矩尺座）、14:00 中五（山案座）、16:30 中六（矩尺座）。星期日矩尺座：12:45 中五、15:15 中四、16:30 中六。不排星期五。出勤日優先矩尺座或山案座。本版星期六中文不避撞科。",
    ),
    (
        "Cyndi Ng",
        "英文科",
        3,
        "星期日",
        "兼職。只限星期日；小組班三班自 10:15 起；另預留一個一對一高中英文時段。",
    ),
    ("Jackson Lau", "英文科（中四至中六）", 2, "星期三、星期六", "兼職。星期三一班、星期六一班。兩班皆高中混級英文（接受中四、中五、中六）。出勤日優先矩尺座或山案座，不排 17D／17E。本版星期六 12:45 改英仙座（矩尺讓 Christine）。"),
    (
        "Judy Chu",
        "生物科（中五級至中六級）",
        3,
        "星期五、星期日",
        "兼職。本版先排中六生物兩班、中五生物一班。連堂最多兩堂，其後必須休息一節。星期日兩班中六因 11:30 僅 17D 有空，12:45 須換 17E。",
    ),
    (
        "Henry Wong",
        "生物科（中四級至中六級）",
        3,
        "星期六",
        "兼職。本版只排星期六連續三班（英仙座）：14:00 中四、15:15 中五、16:30 中六。星期五兩班已移除。可連三堂。星期三尚未確定（預計 8 月 27 日），不排。",
    ),
    (
        "Leo Chan",
        "數學科、物理科",
        5,
        "星期六",
        "兼職。班別時間已確認鎖定。本版五班（數學二、物理三），全部星期六。星期日仍屬問卷已確認可用日，本版移除原有中五物理（B）。星期四、五尚未確定，不排平日。可連續三堂。",
    ),
    (
        "Liam Lai",
        "數學科",
        4,
        "星期六、星期日",
        "兼職。班別時間已確認鎖定。本輪按「3–4」排四班，只教中二、中三數學：星期六兩班、星期日兩班。星期二、三尚未確定，不排該兩日。連堂上限兩堂。12 月中至 1 月頭或外出，屆時按校曆另議。",
    ),
    (
        "Emma Cai",
        "英文科",
        4,
        "星期日",
        "兼職。每周 3–4 班，四班集中星期日一天：中二、中一、中三兩班。不教中六英文。12:45-14:00、14:00-15:15、休息一節、16:30-17:45、17:45-19:00。可連續三堂。",
    ),
    (
        "Cheryl Ng",
        "英文科",
        2,
        "星期六",
        "兼職。只限星期六 10:15 及 11:30（17E）：中二英文、中一英文。可教數學延伸（M2），本版優先補初中英文第二班故不排 M2。連堂最多兩堂。2027 年 6 月 14 日至 7 月 3 日實習，學年末校曆另議。",
    ),
    (
        "Billy Shek",
        "中文科（中二級、中三級）",
        2,
        "星期六",
        "兼職。只教中二、中三中文；營運指定只排 Katie 放假時段，最多兩班。問卷已確定星期六、日，平日（含星期五）尚未確定，本版不排星期五、不排星期日。本版兩班均星期六 17E：16:30 中三、17:45 中二。可連續三堂，本版連兩堂。",
    ),
    (
        "Phoebe Tam",
        "科學科、化學科",
        6,
        "星期六、星期日",
        "兼職。科學、化學各級各一班。問卷已確定星期六、日；星期二尚未確定，本版不排平日。星期日只填 10:15 至 12:45。連堂最多三堂，其後必須休息一節。本版六班：六 17D 連三（10:15 中五化學、11:30 中二科學、12:45 中六化學）→ 休 14:00 → 15:15 中四化學；日 矩尺座連兩（10:15 中三科學、11:30 中一科學）。",
    ),
]

# 本輪調查回覆（2026-08-15 至 20；既有專科老師路徑，無 C 區科目欄）
SURVEY_OVERVIEW = [
    ["老師", "科目（本輪）", "九月開班", "每周堂數", "已確定日子", "本輪已排"],
    ["Judy Chu", "生物（高中）", "願意", "3–4", "星期一、五、日", "3 班（日中六×2、五中五×1）"],
    ["Leo Chan", "數學、物理", "願意", "9 或以上", "星期六、日（全日可）", "5 班（全數星期六；移除星期日中五物理B）"],
    ["Liam Lai", "數學", "願意", "3–4", "星期六、日", "4 班（只中二、中三；六 2、日 2）"],
    ["Emma Cai", "英文", "願意", "3–4", "星期六、日（全日可）", "4 班（集中星期日；不教中六，改中三）"],
    ["Natalie Kwok", "—", "暫不承接", "—", "—", "不排"],
    ["Rafael Ling", "企會財", "願意", "5–6", "無（完全未掌握）", "不排，待補時段"],
    ["Henry Wong", "生物（高中）", "願意", "3–4", "星期五、星期六", "3 班（只排六 14:00 起連三）"],
    ["Cheryl Ng", "英文（可 M2）", "願意", "1–2", "星期六", "2 班（六 10:15 中二、11:30 中一）"],
    ["Billy Shek", "中文（中二、中三）", "願意", "未定", "星期六、日", "2 班（只六；16:30 中三、17:45 中二）"],
    ["Phoebe Tam", "科學、化學", "願意", "未填", "星期六、日", "6 班（科學／化學各級各 1）"],
]

SURVEY_CONSTRAINTS = [
    ["老師", "尚未確定", "連堂", "備註"],
    [
        "Judy Chu",
        "星期三（預計 8 月 18 日）",
        "最多 2 堂；其後必須休息 1 節",
        "平日高中最早 17:45；一 19:00 不可故星期一無法排中五／中六（與現有數學／中文撞級）。日 14:00 起不可。本版星期日上午先排兩班中六生物。",
    ],
    [
        "Leo Chan",
        "星期四、五（預計 9 月 5 日）",
        "最多 3 堂；可連續編排",
        "一至三不可。本版按營運排五班，全部星期六；移除星期日中五物理（B），不排滿「9 或以上」。",
    ],
    [
        "Liam Lai",
        "星期二、三（預計 9 月 1 日）",
        "最多 2 堂；可連續編排",
        "六只用不空白且標「可」之時段。日 14:00 起不可。本版只教中二、中三。12 月中至 1 月頭或外出，屆時按校曆另議。",
    ],
    [
        "Emma Cai",
        "星期一至五（預計 8 月 20 日）",
        "最多 3 堂；可連續編排",
        "平日雖有部分「可」，日子標尚未確定，本版不排平日。四班集中星期日；不教中六英文，改中三英文第二班。",
    ],
    [
        "Natalie Kwok",
        "—",
        "—",
        "想專注學業；私人課程仍會繼續。",
    ],
    [
        "Rafael Ling",
        "星期一至五（預計 8 月 25 日）；週末亦未填時段",
        "最多 3 堂；可連續編排",
        "時段確定後另開一輪，班數 5–6、年級平均分佈。",
    ],
    [
        "Henry Wong",
        "星期三（預計 8 月 27 日）",
        "最多 3 堂；可連續編排",
        "一至四不可。問卷星期五可用，但本版移除星期五兩堂。星期六只用 14:00 至 17:45，連排中四、中五、中六生物。星期日僅 19:00 較不優先，本版不用。",
    ],
    [
        "Cheryl Ng",
        "—",
        "最多 2 堂；可連續編排",
        "只星期六 10:15 及 11:30 可。可教英文與 M2；本版兩班均為初中英文。2027 年 6 月 14 日至 7 月 3 日實習，學年末校曆另議。",
    ],
    [
        "Billy Shek",
        "星期一至五",
        "最多 3 堂；可連續編排",
        "只教中二、中三中文。營運指定只排 Katie 放假時段、最多兩班。星期五尚未確定故不排；星期日雖可，本版不排（Katie 當值）。六 10:15／19:00 較不優先。原 12:45 中一改中二後撞 Leo 中二數學，改排 16:30 中三、17:45 中二（17E）。",
    ],
    [
        "Phoebe Tam",
        "星期二；其餘平日未填",
        "最多 3 堂；其後必須休息 1 節",
        "科學、化學各級各一。六 10:15 至 16:30 可，17:45 起空白。日只填 10:15 至 12:45。本版六四班（17D）、日兩班（矩尺座）。",
    ],
]

SURVEY_SLOT_NOTES = [
    "時段選項：可／較不優先／不可／未確定。空白＝該格未填，本輪視作不可用。",
    "既有專科老師不填科目年級；科目由營運確認：Judy／Henry 生物、Leo 數學與物理、Liam 數學、Emma／Cheryl 英文（Cheryl 可 M2）、Rafael 企會財、Billy 初中中文、Phoebe 科學與化學。",
    "本版以 ver. 3.9 已排格為底，Christine 星期六三堂改中四、中五、中六（不避撞科）；時段與課室不變。已確認老師班別時間鎖不變。",
]

PACKING_SECTIONS = [
    (
        "2.1.1 通則",
        [
            "同日順接只適用星期五、星期六、星期日；星期一至星期四不強制順接。",
            "Judy 先排中六生物兩班、中五生物一班；中四生物由 Henry 承接。",
            "Emma 四班集中星期日，不教中六英文；Cheryl 星期六上午兩班初中英文。",
            "Billy 只排星期六兩班中二／中三中文（Katie 放假時段）；Phoebe 科學、化學各級各一。",
            "Christine 本版只排星期六、星期日；取消星期五。星期六 12:45／16:30 用矩尺（Jackson 12:45 改英仙、Mark 中四數學改 17:45）。",
            "Mark 星期六 12:45 午膳不標示；Jackson 該格用英仙座。",
        ],
    ),
    (
        "2.1.2 中一級",
        [
            "星期六：數學 10:15 → 英文 11:30；另開 15:15 數學（C）。",
            "星期日：中文 10:15 → 科學 11:30 → 英文 14:00；另開 16:30 中文。",
        ],
    ),
    (
        "2.1.3 中二級",
        [
            "星期六：英文 10:15 → 科學 11:30 → 數學 12:45（Leo）→ 14:00 數學（Liam）；17:45 中文（Billy）。",
            "星期日：數學 10:15 → 中文 11:30 → 英文 12:45；另開 15:15 中文。",
        ],
    ),
    (
        "2.1.4 中三級",
        [
            "星期六：數學 10:15（Mark）；15:15 數學（Liam）→ 中文 16:30（Billy）。",
            "星期日：科學 10:15 → 數學 11:30；中文 14:00 → 英文 16:30 → 英文 17:45（平行班）。",
        ],
    ),
    (
        "2.1.5 中四級",
        [
            "星期六：物理 11:30 → 中文／英文 12:45（同時）→ 生物 14:00 → 化學 15:15；另有數學。",
            "星期日：中文 15:15。",
        ],
    ),
    (
        "2.1.6 中五級",
        [
            "星期五：生物 17:45。",
            "星期六：化學 10:15 → 數學 11:30（D）；中文 14:00；下午 生物 15:15 → 物理 16:30。",
            "星期日：中文 12:45。",
        ],
    ),
    (
        "2.1.7 中六級",
        [
            "星期六：化學 12:45 → 數學 14:00（D）→ 物理 15:15 → 生物／中文 16:30（同時）。",
            "星期日：英文 10:15 → 生物 11:30（A）；生物 12:45（B，平行班）→ 中文 16:30。",
        ],
    ),
]

PRINCIPLE_SECTIONS = [
    (
        "1.1 學年與格網",
        [
            "適用學年為 2627（2026-09-01 至 2027-06-30），常規專科班，每周固定逢星期與時段。",
            "每節 75 分鐘。",
            "最遲一節為 19:00-20:15。",
            "不排 20:15-21:30。",
            "目前星期六、星期日不排 09:00-10:15。",
            "週末最早一節為 10:15-11:30。",
        ],
    ),
    (
        "1.2 課室",
        [
            "可用課室為 17D、17E、矩尺座、英仙座、山案座；17K 停用。",
            "平日 16:30 起 17D 列作功課輔導班專用，常規班不使用 17D。",
            "各天詳表最右欄為該時段空房數；返學時間或不排課之列仍計空房（該列課室皆空則為 5）。",
        ],
    ),
    (
        "1.3 年級與時段",
        [
            "平日 15:15-16:30 或之前為返學時間，不排常規班（詳表淺灰標示）。",
            "平日年級時段：中一級至中三級自 16:30 起；中四級至中六級自 17:45 起。",
            "週末除 09:00 限制外，年級不限最早時段，仍禁止末節。",
        ],
    ),
    (
        "1.4 不重疊與順接",
        [
            "同一老師、同一課室、同年級不同科目，同時段均不可重疊。本版例外：Christine 星期六中文不避撞科（12:45 中四與 Jackson 英文同時；16:30 中六與 Henry 生物同時）。",
            "同日順接只適用星期五、星期六、星期日：同一年級該三日宜有不同科目連續時段順接，避免天地堂。星期一至星期四不強制順接。",
            "本輪按已回覆老師的每周堂數編排；同年級同科班數盡量平均，不以單一年級堆疊。",
            "時間表不出現「待確認老師」。時段尚未掌握者不佔格，另列未排。",
        ],
    ),
    (
        "1.5 老師節奏",
        [
            "連堂完全跟各老師問卷意願，不再統一規定「連兩節後必須空一格」。",
            "同日最多五節。",
            "兼職相鄰堂之間空檔最多一格；僅 Katie 可留較大空檔。",
            "每周堂數不多者，能同一天完成則不拆兩天。",
            "同一老師同一出勤日，班別盡量安排於同一課室。",
            "不預留三人開會空檔。",
        ],
    ),
    (
        "1.6 本版老師約束",
        [
            "Mark Yu 出勤日優先矩尺座；不排星期三、星期五、星期日。平日每日最多三班，週末每日最多五班。星期六 12:45-14:00 午膳（時間表不標示）；午後 14:00、15:15、17:45（16:30 矩尺讓 Christine）。",
            "Katie 放假星期五、星期六；本版十三班。平日 14:00 至最後一節；週末 09:00-18:00（中間一節食飯休息）。星期一至四每日兩堂（16:30、17:45）；星期日五班。",
            "Christine Fan 出勤日優先矩尺座或山案座；星期日班別不得早於 11:30。本版只出勤星期六、星期日；不排星期五。星期六 12:45 中四、14:00 中五、16:30 中六；本版不避撞科。",
            "Cyndi Ng 星期日小組班自 10:15 開始；同日另預留一個一對一高中英文時段。",
            "Jackson Lau 出勤為星期三一班、星期六一班；優先矩尺座或山案座，不排 17D／17E。本版星期六 12:45 用英仙座。",
            "Liam Lai 班別時間已確認鎖定；本版只教中二、中三數學。",
            "Leo Chan 班別時間已確認鎖定；本版五班（數學二、物理三），全部星期六；移除星期日中五物理（B）。",
            "Emma Cai 本版不教中六英文。",
            "Henry Wong 本版生物三班：只排星期六 14:00 中四、15:15 中五、16:30 中六；英仙座連續三堂。",
            "Cheryl Ng 本版英文兩班：只星期六 10:15 及 11:30；可教 M2 本版不排。",
            "Billy Shek 本版只教中二、中三中文兩班，只星期六（Katie 放假時段）：16:30 中三、17:45 中二（17E）；不排星期五、星期日。",
            "Phoebe Tam 本版科學、化學各級各一，共六班：星期六四班、星期日兩班；連三後必須休息一節。",
        ],
    ),
]

VERSION_DIFFS = [
    "Katie 取消星期一至四 19:00-20:15 四堂。本版十三班：一至四每日兩堂（16:30、17:45）加星期日五堂。",
    "中一／中二／中三中文班號按星期→時段重編為由 A 起連續。高中中文班號不變。",
    "小組班由 70 減至 66。已確認老師班別時間鎖不變。",
]

PENDING_BY_TEACHER = [
    "Rafael Ling：企會財，意願 5–6 班，時段完全未掌握（預計 8 月 25 日），本版不佔格。",
    "Natalie Kwok：暫不承接專科班。",
    "Billy Shek：星期五尚未確定，不排；星期日雖問卷可，營運指定只排 Katie 放假時段，本版不排日。",
    "Phoebe Tam：星期二尚未確定；星期六 17:45 起及星期日 14:00 起空白，本版不用。",
    "Henry Wong：本版只排星期六連續 3 班；星期三尚未確定（預計 8 月 27 日），不佔格。",
    "Leo Chan：現行五班（全數星期六）班別時間已確認鎖定；意願 9 或以上，星期四、五尚未確定，不另加班。",
    "Liam Lai：現行四班班別時間已確認鎖定；星期二、三尚未確定；12 月中至 1 月頭或外出。",
    "Emma Cai：平日尚未確定日子。",
    "Judy Chu：意願 3–4 班，本版 3 班；星期一無法排高中生物（撞級）。",
    "Cheryl Ng：可教 M2 本版不排；2027 年 6 月 14 日至 7 月 3 日實習，學年末校曆另議。",
    "Christine Fan：本版中四／中五／中六中文各 2 班。",
    "Mark Yu：週末每日最多 5 班，本版星期六已滿 5。",
]

PENDING_BY_SUBJECT = [
    "英文科中六只 1 班（Emma 改中三），欠第二班。",
    "英文科中一至中三本版已各 2 班。",
    "企會財尚未開班。",
    "數學延伸（M2）本版不排。",
    "初中科學、高中化學本版已各級 1 班（Phoebe）。",
    "中五級物理科移除星期日 B 班後，本版只餘 1 班。",
    "中四級生物科移除星期五班後，本版只餘 1 班。",
    "中文各級本版已達每級 ≥2；數學各級已達每級 ≥2。",
]


def teacher_hours_text(n_classes: int) -> str:
    total_min = n_classes * MINUTES_PER_CLASS
    hours, minutes = divmod(total_min, 60)
    if minutes:
        return f"{n_classes} 班 × {MINUTES_PER_CLASS} 分鐘＝每周授課 {hours} 小時 {minutes} 分鐘（合共 {total_min} 分鐘）"
    return f"{n_classes} 班 × {MINUTES_PER_CLASS} 分鐘＝每周授課 {hours} 小時（合共 {total_min} 分鐘）"



def set_run_font(run, size_pt: float = 12, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME_EA
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME_EA)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME_EA)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_EA)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = False
    run.font.color.rgb = color if color is not None else RGBColor(0, 0, 0)


# 老師周視圖：指定不排日／放假（灰底，無斜線）
STAFF_OFF_DAYS = {
    "Katie": {4, 5},  # 星期五、星期六放假
    "Mark Yu": {2, 4, 6},  # 不排星期三、星期五、星期日
    "Christine Fan": {0, 1, 2, 3, 4},  # 不排星期一至五
}
STAFF_OFF_LABEL = {
    "Katie": "放假",
    "Mark Yu": "不排",
    "Christine Fan": "不排",
}
KATIE_OFF_DAYS = STAFF_OFF_DAYS["Katie"]
KATIE_WEEKDAY_DUTY_FIRST_SLOT = 4  # 14:00
KATIE_WEEKEND_DUTY_LAST_SLOT = 6  # 16:30–17:45（約至 18:00）
KATIE_WEEKEND_LUNCH_SLOT = 3  # 12:45–14:00
COLOR_MUTED = RGBColor(0x80, 0x80, 0x80)
FILL_FREE = "EDEDED"
FILL_OFF = "C0C0C0"


def set_style_font(style, size_pt: float = 12, *, bold: bool = False) -> None:
    style.font.name = FONT_NAME_EA
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = False
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_NAME_EA)
    rfonts.set(qn("w:hAnsi"), FONT_NAME_EA)
    rfonts.set(qn("w:eastAsia"), FONT_NAME_EA)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)
    if bold:
        rpr.append(OxmlElement("w:b"))
        rpr.append(OxmlElement("w:bCs"))


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_para(doc: Document, text: str, *, size: float = 12, bold: bool = False, space_after: float = 6, keep_with_next: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.keep_with_next = keep_with_next
    run = p.add_run(text)
    set_run_font(run, size, bold)


def add_bullets(
    doc: Document,
    lines: list[str],
    *,
    size: float = 12,
    space_after: float = 4,
    keep_with_next: bool = False,
) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.keep_with_next = keep_with_next
        run = p.add_run(f"• {line}")
        set_run_font(run, size)


def prevent_row_split(table) -> None:
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def add_heading(
    doc: Document,
    text: str,
    level: int,
    *,
    page_break: bool = False,
    page_break_before: bool = False,
    keep_with_next: bool = False,
    space_after: float | None = None,
) -> None:
    """Word built-in Heading 1–3 so TOC / STYLEREF work."""
    if page_break:
        add_page_break(doc)
    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    size = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}[level]
    p = doc.add_paragraph(text, style=style_name)
    p.paragraph_format.space_before = Pt(0 if level == 1 else 12)
    default_after = 12 if level == 1 else 8
    p.paragraph_format.space_after = Pt(default_after if space_after is None else space_after)
    p.paragraph_format.keep_with_next = keep_with_next
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs:
        set_run_font(run, size, bold=True)
        rpr = run._element.get_or_add_rPr()
        rpr.append(OxmlElement("w:b"))
        rpr.append(OxmlElement("w:bCs"))


def add_chapter(doc: Document, text: str, *, first: bool = False) -> None:
    add_heading(doc, text, 1, page_break=not first)


def add_section_title(doc: Document, text: str, *, keep_with_next: bool = False) -> None:
    add_heading(doc, text, 2, keep_with_next=keep_with_next)


SCHEME_EVEN_HEADER = f"明學教育 2627學年常規時間表方案（ver {VERSION}）"
TEACHERS_EVEN_HEADER = f"明學教育 2627學年各老師時間表（ver {VERSION}）"
WEEKLY_EVEN_HEADER = f"明學教育 2627學年周時間表（ver {VERSION}）"


def add_field_with_placeholder(paragraph, instr: str, placeholder: str = " ", *, size: float = 10) -> None:
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(begin)
    r.append(instr_el)
    r.append(sep)
    r.append(t)
    r.append(end)
    set_run_font(run, size)


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_field_with_placeholder(p, r' TOC \o "1-3" \h \z \u ', "（請於 Word 更新目錄）", size=12)


def clear_paragraph(p) -> None:
    for child in list(p._p):
        if child.tag.endswith("}r"):
            p._p.remove(child)


def write_page_footer_para(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    r = paragraph.add_run("第")
    set_run_font(r, 10)
    add_field_with_placeholder(paragraph, " PAGE ", "1")
    r = paragraph.add_run("頁（共")
    set_run_font(r, 10)
    add_field_with_placeholder(paragraph, " NUMPAGES ", "1")
    r = paragraph.add_run("頁）")
    set_run_font(r, 10)


def write_header_para(paragraph, text: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, 10, color=COLOR_MUTED)


def write_styleref_header(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    clear_paragraph(paragraph)
    # Traditional Chinese Word uses 標題 1 for Heading 1.
    add_field_with_placeholder(paragraph, ' STYLEREF "標題 1" ', " ")


def enable_even_odd_headers(doc: Document) -> None:
    el = doc.settings.element
    if el.find(qn("w:evenAndOddHeaders")) is None:
        el.append(OxmlElement("w:evenAndOddHeaders"))


def apply_normal_page(section, *, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(A4_H_CM)
        section.page_height = Cm(A4_W_CM)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(A4_W_CM)
        section.page_height = Cm(A4_H_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.header_distance = Cm(HEADER_DISTANCE_CM)
    section.footer_distance = Cm(FOOTER_DISTANCE_CM)


def apply_section_chrome(section, even_text: str, *, first_empty: bool) -> None:
    section.different_first_page_header_footer = first_empty
    section.header.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False
    write_styleref_header(section.header.paragraphs[0])
    write_header_para(section.even_page_header.paragraphs[0], even_text)
    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        clear_paragraph(footer.paragraphs[0])
        write_page_footer_para(footer.paragraphs[0])
    if first_empty:
        clear_paragraph(section.first_page_header.paragraphs[0])


def set_row_height_exact(row, cm_h: float) -> None:
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(Cm(cm_h).twips)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    total = int(sum(Cm(w).twips for w in widths_cm))
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for i, w in enumerate(widths_cm):
            if i < len(grid):
                grid[i].set(qn("w:w"), str(int(Cm(w).twips)))
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def add_docx_data_table(doc: Document, rows: list[list[str]], *, header=True, footer_row=False, size: float = 10):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    last_row = len(rows) - 1
    last_col = len(rows[0]) - 1
    for i, row_vals in enumerate(rows):
        for j, val in enumerate(row_vals):
            is_head = header and i == 0
            is_total = footer_row and (i == last_row or j == last_col)
            write_cell(t.rows[i].cells[j], val, bold=(is_head or is_total), size=size)
            if is_head or (footer_row and i == last_row):
                shade_cell(t.rows[i].cells[j])
    prevent_row_split(t)
    return t


def staff_appendix_table() -> list[list[str]]:
    data = [["老師", "班數", "科目", "出勤日", "備註"]]
    for name, subject, n, days, note in STAFF:
        data.append([name, str(n), subject, days, note])
    return data


def set_cell_nowrap(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    if tcPr.find(qn("w:noWrap")) is None:
        tcPr.append(OxmlElement("w:noWrap"))


def shade_cell(cell, fill: str = "F0F0F0", *, stripe: bool = False, stripe_color: str = "666666") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if stripe:
        shd.set(qn("w:val"), "thinDiagStripe")
        shd.set(qn("w:color"), stripe_color)
    else:
        shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_valign(cell, val: str = "center") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:vAlign"))
    if existing is not None:
        tcPr.remove(existing)
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), val)
    tcPr.append(v_align)


def write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 12,
    color: RGBColor | None = None,
) -> None:
    cell.text = ""
    lines = str(text).split("\n") or [""]
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line)
        set_run_font(run, size, bold, color=color)


CORE_SUBJECTS = ["中文", "英文", "數學"]
ELECTIVE_SUBJECTS = ["科學", "物理", "化學", "生物"]


def count_matrix() -> dict[str, dict[str, int]]:
    subjects = CORE_SUBJECTS + ELECTIVE_SUBJECTS
    m = {s: {g: 0 for g in GRADES} for s in subjects}
    for _, _, _, subj, grade, _, _ in CLASSES:
        if subj not in m:
            m[subj] = {g: 0 for g in GRADES}
        m[subj][grade] += 1
    return m


def count_matrix_table() -> list[list[str]]:
    """各級各科班數：最右欄各科合計、最下一行各級合計。"""
    matrix = count_matrix()
    subjects_order = CORE_SUBJECTS + ELECTIVE_SUBJECTS
    rows = [["科目", "中一級", "中二級", "中三級", "中四級", "中五級", "中六級", "合計"]]
    grade_totals = {g: 0 for g in GRADES}
    for subj in subjects_order:
        counts = [matrix[subj][g] for g in GRADES]
        for g, n in zip(GRADES, counts):
            grade_totals[g] += n
        rows.append([SUBJECT_FULL[subj]] + [str(n) for n in counts] + [str(sum(counts))])
    grand = sum(grade_totals[g] for g in GRADES)
    rows.append(["合計"] + [str(grade_totals[g]) for g in GRADES] + [str(grand)])
    return rows


def subject_grade_class_tables() -> list[tuple[str, list[tuple[str, list[list[str]]]]]]:
    """按科目、其後按年級列出全部小組班。空年級略過。"""
    grouped: dict[str, dict[str, list[tuple[int, int, str, str, str]]]] = defaultdict(
        lambda: defaultdict(list)
    )
    for d, slot, room, subj, grade, teacher, code in CLASSES:
        grouped[subj][grade].append((d, slot, room, teacher, code))
    out: list[tuple[str, list[tuple[str, list[list[str]]]]]] = []
    for subj in CORE_SUBJECTS + ELECTIVE_SUBJECTS:
        grade_blocks: list[tuple[str, list[list[str]]]] = []
        for grade in GRADES:
            items = grouped.get(subj, {}).get(grade, [])
            if not items:
                continue
            items.sort(key=lambda x: (class_section_letter(x[4]), x[0], x[1]))
            rows = [["班別", "任教老師", "星期", "時段", "課室"]]
            for d, slot, room, teacher, code in items:
                rows.append(
                    [
                        class_title(grade, subj, code, teacher),
                        teacher,
                        DAYS[d],
                        SLOT_TIMES[slot],
                        room,
                    ]
                )
            grade_blocks.append((GRADE_FULL[grade], rows))
        if grade_blocks:
            out.append((SUBJECT_FULL[subj], grade_blocks))
    return out


def weekly_summary_rows() -> list[list[str]]:
    rows = [["星期", "時段", "課室", "班別", "任教老師"]]
    items: list[tuple[int, int, str, str, str]] = []
    for d, slot, room, subj, grade, teacher, code in CLASSES:
        items.append((d, slot, room, class_title(grade, subj, code, teacher), teacher))
    for d, slot, room, teacher, title in RESERVED:
        items.append((d, slot, room, title, teacher))
    for d, slot, room, title, teacher in sorted(items, key=lambda x: (x[0], x[1], x[2])):
        rows.append([DAYS[d], SLOT_TIMES[slot], room, title, teacher])
    return rows


WEEKDAY_SCHOOL_LAST_SLOT = 5  # 15:15-16:30 及之前
FILL_SCHOOL = "EDEDED"
MUTED_LABELS = {"返學時間", "不排課"}


def cover_notes() -> list[str]:
    return [
        "本文件為規劃方案紀錄，供營運審閱",
        "內容以排課規則為準，尚未寫入正式班別與排程",
        f"本版為 ver. {VERSION}；ver. {PREV_VERSION} 檔案保留不動",
    ]


def day_notes(day_idx: int) -> list[str]:
    if day_idx >= 5:
        notes = [
            "週末",
            "五室皆可排常規班",
            "最遲 19:00-20:15",
            "本日不排 09:00-10:15",
        ]
    else:
        notes = [
            "平日",
            "15:15-16:30 或之前為返學時間（淺灰）",
            "16:30 起 17D 為功課輔導班專用",
            "最遲 19:00-20:15",
        ]
    if day_idx == 4:
        notes.append("本版不預留開會空檔")
    if day_idx == 5:
        notes.append("Cheryl Ng：中二級英文科（B) 10:15、中一級英文科（B) 11:30，17E")
        notes.append("Henry Wong：中四級生物科（A) 14:00、中五級生物科（B) 15:15、中六級生物科（C) 16:30，英仙座連續三堂")
        notes.append("Christine Fan：中四級中文科 12:45 矩尺座、中五級中文科 14:00 山案座、中六級中文科 16:30 矩尺座（本版不避撞科）")
        notes.append("Jackson Lau：高中英文科（中四至中六）（B) 12:45 改英仙座")
    if day_idx == 6:
        notes.append("Christine Fan 本版自 12:45 起：中五、中四、中六（矩尺座）")
        notes.append("Cyndi Ng 小組班自 10:15 起，並預留一個一對一高中英文時段")
    return notes


def day_grid(day_idx: int) -> list[list[str]]:
    header = ["時段"] + ROOMS + ["空房"]
    grid = [header]
    is_weekday = day_idx <= 4
    is_weekend = day_idx >= 5
    for s, label in enumerate(SLOT_TIMES):
        row = [label]
        school_hours = is_weekday and s <= WEEKDAY_SCHOOL_LAST_SLOT
        empty = 0
        for room in ROOMS:
            if is_weekend and s == 0:
                row.append("不排課")
                empty += 1
                continue
            if school_hours:
                row.append("返學時間")
                empty += 1
                continue
            if is_weekday and room == "17D":
                cell = "功課輔導班"
            else:
                cell = ""
            hits = [
                class_cell_text(grade, subj, teacher, code, slot)
                for d, slot, r, subj, grade, teacher, code in CLASSES
                if d == day_idx and slot == s and r == room
            ]
            reserved_hits = [
                f"{title}\n{teacher}\n{SLOT_TIMES[slot]}"
                for d, slot, r, teacher, title in RESERVED
                if d == day_idx and slot == s and r == room
            ]
            hits.extend(reserved_hits)
            if hits:
                extra = "\n\n".join(hits)
                cell = extra if not cell else f"{cell}\n{extra}"
            if cell:
                row.append(cell)
            else:
                empty += 1
                row.append("—")
        row.append(str(empty))
        grid.append(row)
    return grid


def configure_docx_styles(doc: Document, *, landscape: bool = False) -> None:
    for section in doc.sections:
        apply_normal_page(section, landscape=landscape)
    set_style_font(doc.styles["Normal"], 12)
    for name, size, bold in (
        ("Heading 1", H1_SIZE, True),
        ("Heading 2", H2_SIZE, True),
        ("Heading 3", H3_SIZE, True),
        ("Title", 12, True),
    ):
        if name in doc.styles:
            set_style_font(doc.styles[name], size, bold=bold)


def build_docx(path: Path) -> None:
    doc = Document()
    configure_docx_styles(doc)
    enable_even_odd_headers(doc)
    first = doc.sections[0]
    apply_section_chrome(first, SCHEME_EVEN_HEADER, first_empty=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("明學教育")
    set_run_font(run, 12, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("2627 學年常規專科班時間表")
    set_run_font(run, 12, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"方案紀錄 ver. {VERSION}｜對照 ver. {PREV_VERSION}｜未定稿入庫")
    set_run_font(run, 12)

    add_bullets(doc, cover_notes())
    add_para(doc, f"與 ver. {PREV_VERSION} 之分別", size=12, bold=True, space_after=6)
    add_bullets(doc, VERSION_DIFFS)

    add_page_break(doc)
    add_para(doc, "目錄", size=H1_SIZE, bold=True, space_after=8)
    add_toc_field(doc)

    add_chapter(doc, "1. 排程原則")
    for heading, lines in PRINCIPLE_SECTIONS:
        add_section_title(doc, heading, keep_with_next=True)
        add_bullets(doc, lines)

    add_chapter(doc, "2. 各級開科情況")
    add_section_title(doc, "2.1 各年級順接方案")
    for heading, lines in PACKING_SECTIONS:
        add_heading(doc, heading, 3, keep_with_next=True)
        add_bullets(doc, lines)
    add_section_title(doc, "2.2 各級各科班數")
    add_bullets(
        doc,
        [
            "下表為各級各科已排小組班數",
            "最右欄為各科合計",
            "最下一行為各級合計",
            "不含一對一預留",
        ],
    )
    add_docx_data_table(doc, count_matrix_table(), footer_row=True, size=12)

    add_chapter(doc, "5. 各科各級班別列表")
    add_bullets(doc, ["以下按科目、其後按年級列出全部已排小組班", "一對一預留見章末"])
    n_subj = 0
    for si, (subj_full, grade_blocks) in enumerate(subject_grade_class_tables(), start=1):
        n_subj = si
        add_section_title(doc, f"5.{si} {subj_full}")
        for gi, (grade_full, rows) in enumerate(grade_blocks, start=1):
            add_heading(doc, f"5.{si}.{gi} {grade_full}", 3, keep_with_next=True)
            add_docx_data_table(doc, rows, size=10)
    if RESERVED:
        add_section_title(doc, f"5.{n_subj + 1} 預留時段（不計小組班）")
        for d, slot, room, teacher, title_t in RESERVED:
            add_para(
                doc,
                f"{title_t}｜{teacher}｜{DAYS[d]} {SLOT_TIMES[slot]}｜{room}",
                size=12,
                space_after=4,
            )

    add_chapter(doc, "6. 未排與待排")
    add_bullets(doc, ["本輪按老師已確定檔期與每周堂數編排", "不以「待確認老師」佔用課室格"])
    add_section_title(doc, "6.1 按老師")
    add_bullets(doc, PENDING_BY_TEACHER)
    add_section_title(doc, "6.2 按科目")
    add_bullets(doc, PENDING_BY_SUBJECT)

    add_chapter(doc, "7. 所有班別清單")
    add_bullets(doc, ["以下按星期與時段列出全部已排班及預留時段", f"已排小組班合計 {len(CLASSES)} 班", f"另預留時段 {len(RESERVED)} 個"])
    add_docx_data_table(doc, weekly_summary_rows(), size=12)

    add_chapter(doc, "附表 本輪老師回覆與出勤")
    add_bullets(
        doc,
        [
            "以下整理調查表各人意願，並結合各員工出勤、班數、科目",
            "科目欄為營運確認，非表格 C 區自填（既有專科老師不經 C 區）",
        ],
    )
    add_bullets(doc, SURVEY_SLOT_NOTES)
    add_section_title(doc, "附表.1 調查回覆")
    add_docx_data_table(doc, SURVEY_OVERVIEW, size=10)
    add_section_title(doc, "附表.2 檔期與連堂約束")
    add_docx_data_table(doc, SURVEY_CONSTRAINTS, size=10)
    add_section_title(doc, "附表.3 各員工出勤日、班數、科目")
    add_docx_data_table(doc, staff_appendix_table(), size=10)

    add_para(doc, "— 完 —", size=12)
    doc.save(path)


WEEKLY_HEADER_ROW_CM = 0.72
WEEKLY_DATA_ROW_CM = 1.50


def add_weekly_day_table(doc: Document, day_idx: int, widths: list[float]) -> None:
    grid = day_grid(day_idx)
    td = doc.add_table(rows=len(grid), cols=len(grid[0]))
    td.style = "Table Grid"
    last_col = len(grid[0]) - 1
    for i, row_vals in enumerate(grid):
        set_row_height_exact(td.rows[i], WEEKLY_DATA_ROW_CM if i else WEEKLY_HEADER_ROW_CM)
        for j, val in enumerate(row_vals):
            muted = val in MUTED_LABELS
            write_cell(
                td.rows[i].cells[j],
                val,
                bold=(i == 0 or j == 0 or j == last_col),
                size=8 if i else 10,
                color=COLOR_MUTED if muted else None,
            )
            set_cell_valign(td.rows[i].cells[j])
            if i == 0:
                shade_cell(td.rows[i].cells[j])
            elif muted:
                shade_cell(td.rows[i].cells[j], FILL_SCHOOL)
            if j == 0 or j == last_col:
                for para in td.rows[i].cells[j].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_nowrap(td.rows[i].cells[j])
    set_table_widths(td, widths)
    prevent_row_split(td)


def build_weekly_docx(path: Path) -> None:
    doc = Document()
    configure_docx_styles(doc)
    enable_even_odd_headers(doc)
    apply_section_chrome(doc.sections[0], WEEKLY_EVEN_HEADER, first_empty=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("明學教育")
    set_run_font(run, 12, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("2627 學年常規專科班周時間表")
    set_run_font(run, 12, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"獨立附件｜對應方案紀錄 ver. {VERSION}｜未定稿入庫")
    set_run_font(run, 12)

    add_bullets(
        doc,
        [
            "本文件獨立於方案正文",
            f"已排小組班合計 {len(CLASSES)} 班；另預留時段 {len(RESERVED)} 個",
            "每日一頁、整表不拆開；時間表頁面為橫向",
            "格高固定；無課仍維持同一高度",
            "平日 15:15-16:30 或之前淺灰標返學時間，該列仍計空房",
            "最右欄為該時段空房數",
        ],
    )

    add_page_break(doc)
    add_para(doc, "目錄", size=H1_SIZE, bold=True, space_after=8)
    add_toc_field(doc)

    land = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_normal_page(land, landscape=True)
    apply_section_chrome(land, WEEKLY_EVEN_HEADER, first_empty=False)

    usable = A4_H_CM - MARGIN_CM * 2
    time_w, empty_w = 2.15, 1.2
    room_w = (usable - time_w - empty_w) / 5
    widths = [time_w] + [room_w] * 5 + [empty_w]
    for day_idx, day_name in enumerate(DAYS):
        add_heading(
            doc,
            f"{day_idx + 1}. {day_name}",
            1,
            page_break_before=(day_idx > 0),
            keep_with_next=True,
            space_after=6,
        )
        add_weekly_day_table(doc, day_idx, widths)
    doc.save(path)



def validate() -> None:
    assert len(CLASSES) == 66, len(CLASSES)
    codes = class_codes_rows()
    assert len(codes) == len(CLASSES)
    assert len({row["班別顯示碼（course_code_full）"] for row in codes}) == len(CLASSES)
    for row in codes:
        assert row["課程名稱（course_name）"]
    matrix_rows = count_matrix_table()
    assert int(matrix_rows[-1][-1]) == len(CLASSES)
    listed = sum(len(rows) - 1 for _, blocks in subject_grade_class_tables() for _, rows in blocks)
    assert listed == len(CLASSES)
    teachers = {c[5] for c in CLASSES}
    assert "待確認老師" not in teachers
    assert "Natalie Kwok" not in teachers
    assert "Rafael Ling" not in teachers
    assert "Cyndi" not in teachers  # full name Cyndi Ng
    for teacher, expected in CONFIRMED_TEACHER_CLASS_TIMES.items():
        actual = {
            (day_idx, slot_idx, code)
            for day_idx, slot_idx, _room, _subj, _grade, name, code in CLASSES
            if name == teacher
        }
        assert actual == expected, (teacher, actual, expected)
    t_busy: dict = defaultdict(list)
    r_busy: dict = defaultdict(list)
    g_busy: dict = defaultdict(list)
    for d, s, room, subj, grade, teacher, code in CLASSES:
        t_busy[(d, s)].append(teacher)
        r_busy[(d, s, room)].append(code)
        g_busy[(d, s, grade)].append(subj)
        if d <= 4:
            if grade in {"S1", "S2", "S3"}:
                assert s >= 6, (code, s)
            if grade in {"S4", "S5", "S6"}:
                assert s >= 7, (code, s)
            assert room != "17D", code
        if d >= 5:
            assert s >= 1
        assert s <= 8
        if teacher == "Judy Chu":
            assert subj == "生物" and grade in {"S5", "S6"}
        if teacher == "Emma Cai":
            assert subj == "英文" and d == 6 and grade != "S6"
        if teacher == "Jackson Lau":
            assert room in {"矩尺座", "山案座", "英仙座"}
        if teacher == "Mark Yu":
            assert room == "矩尺座"
        if teacher == "Christine Fan":
            assert room in {"矩尺座", "山案座"}
        if teacher == "Liam Lai":
            assert subj == "數學" and grade in {"S2", "S3"}
        if teacher == "Leo Chan":
            assert subj in {"數學", "物理"}
        if teacher == "Henry Wong":
            assert subj == "生物" and grade in {"S4", "S5", "S6"} and room == "英仙座"
        if teacher == "Cheryl Ng":
            assert subj == "英文" and d == 5 and room == "17E"
        if teacher == "Billy Shek":
            assert subj == "中文" and grade in {"S2", "S3"} and d == 5
        if teacher == "Phoebe Tam":
            assert subj in {"科學", "化學"}
            if subj == "科學":
                assert grade in {"S1", "S2", "S3"}
            else:
                assert grade in {"S4", "S5", "S6"}
    for d, s, room, teacher, title in RESERVED:
        t_busy[(d, s)].append(teacher)
        r_busy[(d, s, room)].append(title)
    for k, v in t_busy.items():
        assert len(v) == len(set(v)), k
    for k, v in r_busy.items():
        assert len(v) == 1, (k, v)
    for k, v in g_busy.items():
        assert len(v) == len(set(v)), (k, v)
    for d, s, _, _, _, teacher, _ in CLASSES:
        if teacher == "Christine Fan" and d == 6:
            assert s >= 2
        if teacher == "Cyndi Ng":
            assert d == 6 and s >= 1
        # Weekend: no 09:00
        if d >= 5:
            assert s >= 1
    assert {c[0] for c in CLASSES if c[5] == "Jackson Lau"} == {2, 5}
    assert {c[0] for c in CLASSES if c[5] == "Katie"} == {0, 1, 2, 3, 6}
    assert {c[0] for c in CLASSES if c[5] == "Mark Yu"} == {0, 1, 3, 5}
    # Mark 一／二連三（slots 6–8）；Katie 一至四每日兩堂（slots 6–7）
    for d in (0, 1):
        slots = sorted(c[1] for c in CLASSES if c[5] == "Mark Yu" and c[0] == d)
        assert slots == [6, 7, 8], ("Mark Yu", d, slots)
    for d in (0, 1, 2, 3):
        slots = sorted(c[1] for c in CLASSES if c[5] == "Katie" and c[0] == d)
        assert slots == [6, 7], ("Katie", d, slots)
    mark_thu = sorted(c[1] for c in CLASSES if c[5] == "Mark Yu" and c[0] == 3)
    assert mark_thu == [7, 8]
    mark_thu_grades = {c[4] for c in CLASSES if c[5] == "Mark Yu" and c[0] == 3}
    assert mark_thu_grades <= {"S4", "S5", "S6"}
    mark_sat = sorted(c[1] for c in CLASSES if c[5] == "Mark Yu" and c[0] == 5)
    assert mark_sat == [1, 2, 4, 5, 7]
    assert {c[4] for c in CLASSES if c[5] == "Mark Yu" and c[0] == 5} == {"S1", "S3", "S4", "S5", "S6"}
    assert 3 not in mark_sat  # 12:45 午膳不排
    mark_per_day: dict[int, int] = defaultdict(int)
    for d, _s, _r, _subj, _g, teacher, _c in CLASSES:
        if teacher == "Mark Yu":
            mark_per_day[d] += 1
    for d, n in mark_per_day.items():
        if d <= 4:
            assert n <= 3, (d, n)
        else:
            assert n <= 5, (d, n)
    assert sum(1 for c in CLASSES if c[5] == "Katie") == 13
    assert sorted(c[1] for c in CLASSES if c[5] == "Katie" and c[0] == 6) == [1, 2, 4, 5, 6]
    assert sorted(c[1] for c in CLASSES if c[5] == "Christine Fan" and c[0] == 6) == [3, 5, 6]
    assert {c[0] for c in CLASSES if c[5] == "Christine Fan"} == {5, 6}
    assert sorted(c[1] for c in CLASSES if c[5] == "Christine Fan" and c[0] == 5) == [3, 4, 6]
    assert {(c[1], c[4], c[2]) for c in CLASSES if c[5] == "Christine Fan" and c[0] == 5} == {
        (3, "S4", "矩尺座"),
        (4, "S5", "山案座"),
        (6, "S6", "矩尺座"),
    }
    assert all(c[2] == "矩尺座" for c in CLASSES if c[5] == "Christine Fan" and c[0] == 6)
    assert not any(c[0] <= 4 and c[5] == "Christine Fan" for c in CLASSES)
    christine_sun = {(c[1], c[4]) for c in CLASSES if c[5] == "Christine Fan" and c[0] == 6}
    assert christine_sun == {(3, "S5"), (5, "S4"), (6, "S6")}
    for d in range(7):
        grid = day_grid(d)
        assert grid[0][-1] == "空房"
        for ri, row in enumerate(grid[1:], start=1):
            assert row[-1].isdigit(), (d, ri, row[-1])
            slot = ri - 1
            if d <= 4 and slot <= WEEKDAY_SCHOOL_LAST_SLOT:
                assert row[1] == "返學時間"
                assert row[-1] == "5"
            if d >= 5 and slot == 0:
                assert row[-1] == "5"
    chi_letters: dict[str, list[str]] = defaultdict(list)
    for _d, _s, _r, subj, grade, _t, code in CLASSES:
        if subj == "中文":
            chi_letters[grade].append(class_section_letter(code))
    for grade, letters in chi_letters.items():
        expected = [chr(ord("A") + i) for i in range(len(letters))]
        assert sorted(letters) == expected, (grade, letters)
    chi = count_matrix()["中文"]
    assert chi["S1"] == 5 and chi["S2"] == 5 and chi["S3"] == 5
    assert chi["S4"] == 2 and chi["S5"] == 2 and chi["S6"] == 2
    assert len(RESERVED) == 1 and RESERVED[0][4].startswith("一對一高中英文科")
    rooms = defaultdict(set)
    for d, _, room, _, _, teacher, _ in CLASSES:
        rooms[(teacher, d)].add(room)
    for d, _, room, teacher, _ in RESERVED:
        rooms[(teacher, d)].add(room)
    for k, rs in rooms.items():
        if k in {("Judy Chu", 6), ("Christine Fan", 5)}:
            continue  # 日 Judy 換房；六 Christine 14:00 矩尺為 Mark 數學故用山案
        assert len(rs) == 1, (k, rs)
    for name, _subject, n, _days, _note in STAFF:
        actual = sum(1 for c in CLASSES if c[5] == name)
        assert actual == n, (name, actual, n)
    allow_three = {
        ("Mark Yu", 0),
        ("Mark Yu", 1),
        ("Katie", 0),
        ("Katie", 1),
        ("Katie", 2),
        ("Katie", 3),
        ("Katie", 6),
        ("Leo Chan", 5),
        ("Henry Wong", 5),
        ("Mark Yu", 5),
        ("Phoebe Tam", 5),
    }
    by_td: dict[tuple[str, int], list[int]] = defaultdict(list)
    for d, s, _r, _subj, _g, teacher, _c in CLASSES:
        by_td[(teacher, d)].append(s)
    for d, s, _r, teacher, _t in RESERVED:
        by_td[(teacher, d)].append(s)
    for (teacher, d), slots in by_td.items():
        slots = sorted(slots)
        assert len(slots) <= 5, (teacher, d, slots)
        if teacher != "Katie":
            for a, b in zip(slots, slots[1:]):
                assert b - a <= 2, (teacher, d, slots)
        run = 1
        max_run = 1
        for a, b in zip(slots, slots[1:]):
            if b == a + 1:
                run += 1
                max_run = max(max_run, run)
            else:
                run = 1
        limit = 3 if (teacher, d) in allow_three else 2
        assert max_run <= limit, (teacher, d, slots, max_run)
    assert sorted(c[1] for c in CLASSES if c[5] == "Leo Chan" and c[0] == 5) == [1, 2, 3, 5, 6]
    assert not any(c[0] == 6 and c[5] == "Leo Chan" for c in CLASSES)
    assert sum(1 for c in CLASSES if c[5] == "Leo Chan") == 5
    assert sorted(c[1] for c in CLASSES if c[5] == "Liam Lai" and c[0] == 5) == [4, 5]
    assert {c[4] for c in CLASSES if c[5] == "Liam Lai" and c[0] == 5} == {"S2", "S3"}
    assert sorted(c[1] for c in CLASSES if c[5] == "Liam Lai" and c[0] == 6) == [1, 2]
    assert {c[4] for c in CLASSES if c[5] == "Liam Lai"} <= {"S2", "S3"}
    assert sorted(c[1] for c in CLASSES if c[5] == "Emma Cai" and c[0] == 6) == [3, 4, 6, 7]
    assert {c[4] for c in CLASSES if c[5] == "Emma Cai"} == {"S1", "S2", "S3"}
    assert not any(c[0] == 5 and c[5] == "Emma Cai" for c in CLASSES)
    judy_grades = sorted(c[4] for c in CLASSES if c[5] == "Judy Chu")
    assert judy_grades == ["S5", "S6", "S6"]
    assert {c[0] for c in CLASSES if c[5] == "Henry Wong"} == {5}
    assert sorted(c[1] for c in CLASSES if c[5] == "Henry Wong") == [4, 5, 6]
    assert {c[4] for c in CLASSES if c[5] == "Henry Wong"} == {"S4", "S5", "S6"}
    assert {c[6] for c in CLASSES if c[5] == "Henry Wong"} == {"S4生A", "S5生B", "S6生C"}
    assert {c[0] for c in CLASSES if c[5] == "Cheryl Ng"} == {5}
    assert sorted(c[1] for c in CLASSES if c[5] == "Cheryl Ng") == [1, 2]
    assert {c[4] for c in CLASSES if c[5] == "Cheryl Ng"} == {"S1", "S2"}
    assert {c[0] for c in CLASSES if c[5] == "Billy Shek"} == {5}
    assert sorted(c[1] for c in CLASSES if c[5] == "Billy Shek") == [6, 7]
    assert {c[4] for c in CLASSES if c[5] == "Billy Shek"} == {"S2", "S3"}
    assert all(c[2] == "17E" for c in CLASSES if c[5] == "Billy Shek")
    assert sum(1 for c in CLASSES if c[5] == "Phoebe Tam") == 6
    assert sorted(c[1] for c in CLASSES if c[5] == "Phoebe Tam" and c[0] == 5) == [1, 2, 3, 5]
    assert all(c[2] == "17D" for c in CLASSES if c[5] == "Phoebe Tam" and c[0] == 5)
    assert sorted(c[1] for c in CLASSES if c[5] == "Phoebe Tam" and c[0] == 6) == [1, 2]
    assert all(c[2] == "矩尺座" for c in CLASSES if c[5] == "Phoebe Tam" and c[0] == 6)
    sci = count_matrix()["科學"]
    assert sci["S1"] == 1 and sci["S2"] == 1 and sci["S3"] == 1
    chem = count_matrix()["化學"]
    assert chem["S4"] == 1 and chem["S5"] == 1 and chem["S6"] == 1
    eng = count_matrix()["英文"]
    assert eng["S1"] == 2 and eng["S2"] == 2 and eng["S3"] == 2 and eng["S6"] == 1
    bio = count_matrix()["生物"]
    assert bio["S4"] == 1 and bio["S5"] == 2 and bio["S6"] == 3
    physics = count_matrix()["物理"]
    assert physics["S5"] == 1
    katie_grid = teacher_week_grid("Katie")
    assert katie_grid[1][5] == ("放假", "off")
    assert katie_grid[1][6] == ("放假", "off")
    assert katie_grid[4][7] == ("食飯休息", "lunch")
    assert katie_grid[5][1] == ("空堂", "free")  # 星期一 14:00
    assert katie_grid[9][1] == ("空堂", "free")  # 星期一 19:00 本版不排
    mark_grid = teacher_week_grid("Mark Yu")
    assert mark_grid[1][3] == ("不排", "off")  # 星期三
    assert mark_grid[1][5] == ("不排", "off")  # 星期五
    assert mark_grid[1][7] == ("不排", "off")  # 星期日
    christine_grid = teacher_week_grid("Christine Fan")
    assert christine_grid[1][5] == ("不排", "off")  # 星期五
    assert "中四級中文科" in christine_grid[4][6][0]  # 星期六 12:45
    assert "中六級中文科" in christine_grid[7][6][0]  # 星期六 16:30
    scheme_md = build_scheme_md()
    weekly_md = build_weekly_md()
    assert "## 8. 周時間表" not in scheme_md
    assert all(f"## {day_idx + 1}. {day_name}" in weekly_md for day_idx, day_name in enumerate(DAYS))
    for day_idx, day_name in enumerate(DAYS):
        heading = f"## {day_idx + 1}. {day_name}\n"
        rest = weekly_md.split(heading, 1)[1].lstrip()
        assert rest.startswith("| 時段"), rest[:40]
    mon = day_grid(0)
    assert "16:30-17:45" in mon[7][3]
    assert "Mark Yu" in mon[7][3]
    print("validation ok", len(CLASSES), "classes", len(RESERVED), "reserved")


def teacher_week_rows(teacher: str) -> list[list[str]]:
    """Rows for one teacher: 星期, 時段, 課室, 班別."""
    rows = [["星期", "時段", "課室", "班別"]]
    items: list[tuple[int, int, str, str]] = []
    for d, slot, room, subj, grade, name, code in CLASSES:
        if name == teacher:
            items.append((d, slot, room, class_title(grade, subj, code, name)))
    for d, slot, room, name, title in RESERVED:
        if name == teacher:
            items.append((d, slot, room, title))
    for d, slot, room, title in sorted(items, key=lambda x: (x[0], x[1], x[2])):
        rows.append([DAYS[d], SLOT_TIMES[slot], room, title])
    return rows


def katie_week_kind(day: int, slot: int, has_class: bool) -> str:
    """Katie 周視圖格類型：off / class / lunch / free / empty。"""
    if day in KATIE_OFF_DAYS:
        return "off"
    if has_class:
        return "class"
    if day <= 4:
        return "free" if slot >= KATIE_WEEKDAY_DUTY_FIRST_SLOT else "empty"
    if slot == KATIE_WEEKEND_LUNCH_SLOT:
        return "lunch"
    if slot <= KATIE_WEEKEND_DUTY_LAST_SLOT:
        return "free"
    return "empty"


def teacher_week_grid(teacher: str) -> list[list[tuple[str, str]]]:
    """University-style week view: rows = slots, columns = days.

    Each cell is (text, kind) where kind is header / time / class / empty /
    free / lunch / off / blocked.
    """
    cells: dict[tuple[int, int], str] = {}
    for d, slot, room, subj, grade, name, code in CLASSES:
        if name == teacher:
            cells[(d, slot)] = f"{class_title(grade, subj, code, name)}\n{room}"
    for d, slot, room, name, title in RESERVED:
        if name == teacher:
            cells[(d, slot)] = f"{title}\n{room}"

    off_days = STAFF_OFF_DAYS.get(teacher, set())
    off_label = STAFF_OFF_LABEL.get(teacher, "不排")
    header: list[tuple[str, str]] = [("時段", "header")]
    for di, day in enumerate(DAYS):
        header.append((day, "off" if di in off_days else "header"))
    grid: list[list[tuple[str, str]]] = [header]
    for s, time_label in enumerate(SLOT_TIMES):
        row: list[tuple[str, str]] = [(time_label, "time")]
        for d in range(7):
            if d in off_days:
                row.append((off_label, "off"))
                continue
            text = cells.get((d, s))
            if teacher == "Katie":
                kind = katie_week_kind(d, s, text is not None)
                if kind == "class":
                    row.append((text or "", "class"))
                elif kind == "lunch":
                    row.append(("食飯休息", "lunch"))
                elif kind == "free":
                    row.append(("空堂", "free"))
                else:
                    row.append(("—", "empty"))
            elif d >= 5 and s == 0:
                row.append(("不排課", "blocked"))
            elif text:
                row.append((text, "class"))
            else:
                row.append(("—", "empty"))
        grid.append(row)
    return grid


def week_view_legend(name: str) -> str | None:
    if name == "Katie":
        return "圖例：淺灰「空堂」＝當值但非上堂；「食飯休息」＝週末中間一節；灰底「放假」＝全日放假。"
    if name in STAFF_OFF_DAYS:
        return "圖例：灰底「不排」＝指定不排日全日。"
    return None


def apply_week_cell_docx(cell, text: str, kind: str, *, header_like: bool) -> None:
    muted = kind in {"free", "lunch", "off"}
    write_cell(
        cell,
        text,
        bold=header_like or kind == "off",
        size=8,
        color=COLOR_MUTED if muted else None,
    )
    if kind in {"free", "lunch", "off"}:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if kind in {"header", "time"}:
        shade_cell(cell)
    elif kind in {"free", "lunch"}:
        shade_cell(cell, FILL_FREE)
    elif kind == "off":
        shade_cell(cell, FILL_OFF)


def build_teachers_docx(path: Path) -> None:
    doc = Document()
    configure_docx_styles(doc)
    enable_even_odd_headers(doc)
    apply_section_chrome(doc.sections[0], TEACHERS_EVEN_HEADER, first_empty=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("明學教育")
    set_run_font(run, 12, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("2627 學年各老師一周排程")
    set_run_font(run, 12, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"獨立附件｜對應方案紀錄 ver. {VERSION}｜未定稿入庫")
    set_run_font(run, 12)

    add_bullets(
        doc,
        [
            "本文件獨立於全校課室時間表，只按老師列出一周職務",
            "每位老師先列一周總覽，周視圖另頁（橫軸星期、縱軸時段）",
            "班別名稱與時段寫法與方案紀錄一致",
            "指定不排日（Katie 放假五／六；Mark 不排三／五／日；Christine 不排一至五）於周視圖以灰底標示",
            "Katie 另標空堂與食飯休息",
        ],
    )
    add_page_break(doc)
    add_para(doc, "目錄", size=H1_SIZE, bold=True, space_after=8)
    add_toc_field(doc)

    for name, subject, n, days, note in STAFF:
        add_chapter(doc, name, first=False)
        reserved_n = sum(1 for _d, _s, _r, teacher, _t in RESERVED if teacher == name)
        load = f"小組班：{n} 班"
        if reserved_n:
            load += f"；預留時段：{reserved_n} 個"
        add_bullets(
            doc,
            [
                f"科目：{subject}",
                load,
                teacher_hours_text(n),
                f"出勤日：{days}",
                note,
            ],
        )

        add_para(doc, "一周總覽", size=12, bold=True, space_after=6, keep_with_next=True)
        overview = teacher_week_rows(name)
        t = doc.add_table(rows=len(overview), cols=4)
        t.style = "Table Grid"
        for ri, row_vals in enumerate(overview):
            for j, val in enumerate(row_vals):
                write_cell(t.rows[ri].cells[j], val, bold=(ri == 0), size=12)
                if ri == 0:
                    shade_cell(t.rows[ri].cells[j])
        prevent_row_split(t)

        add_page_break(doc)
        add_para(doc, f"{name}｜周視圖", size=12, bold=True, space_after=6, keep_with_next=True)
        add_para(doc, f"科目：{subject}；{teacher_hours_text(n)}", size=12, space_after=8, keep_with_next=True)
        legend = week_view_legend(name)
        if legend:
            add_para(doc, legend, size=12, space_after=8, keep_with_next=True)
        grid = teacher_week_grid(name)
        tg = doc.add_table(rows=len(grid), cols=len(grid[0]))
        tg.style = "Table Grid"
        for ri, row_vals in enumerate(grid):
            for j, (val, kind) in enumerate(row_vals):
                apply_week_cell_docx(tg.rows[ri].cells[j], val, kind, header_like=(ri == 0 or j == 0))
        prevent_row_split(tg)

    add_para(doc, "— 完 —", size=12)
    doc.save(path)



def md_cell(text: object) -> str:
    return str(text).replace("\n", "／").replace("|", "\\|")


def md_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    head = rows[0]
    lines = [
        "| " + " | ".join(md_cell(c) for c in head) + " |",
        "| " + " | ".join("---" for _ in head) + " |",
    ]
    for row in rows[1:]:
        padded = list(row) + [""] * (len(head) - len(row))
        lines.append("| " + " | ".join(md_cell(c) for c in padded[: len(head)]) + " |")
    return "\n".join(lines)


def md_bullets(lines: list[str]) -> str:
    return "\n".join(f"- {line}" for line in lines)


def build_scheme_md() -> str:
    parts: list[str] = [
        "# 明學教育 2627 學年常規專科班時間表",
        "",
        f"方案紀錄 **ver. {VERSION}**｜對照 ver. {PREV_VERSION}｜未定稿入庫",
        "",
        "> 本檔為審閱稿。docx／pdf 只在營運決定出檔後執行 `python3 scripts/generate_2627_timetable_doc.py --word`（Word 內建目錄／頁首頁尾，PDF 由 Word 另存）。",
        "",
        "## 封面說明",
        md_bullets(cover_notes()),
        "",
        f"## 與 ver. {PREV_VERSION} 之分別",
        md_bullets(VERSION_DIFFS),
        "",
        "## 1. 排程原則",
    ]
    for heading, lines in PRINCIPLE_SECTIONS:
        parts.extend(["", f"### {heading}", md_bullets(lines)])
    parts.extend(["", "## 2. 各級開科情況", "", "### 2.1 各年級順接方案"])
    for heading, lines in PACKING_SECTIONS:
        parts.extend(["", f"#### {heading}", md_bullets(lines)])
    parts.extend(
        [
            "",
            "### 2.2 各級各科班數",
            md_bullets(
                [
                    "下表為各級各科已排小組班數",
                    "最右欄為各科合計",
                    "最下一行為各級合計",
                    "不含一對一預留",
                ]
            ),
            "",
            md_table(count_matrix_table()),
            "",
            "## 5. 各科各級班別列表",
            md_bullets(["以下按科目、其後按年級列出全部已排小組班", "一對一預留見章末"]),
        ]
    )
    n_subj = 0
    for si, (subj_full, grade_blocks) in enumerate(subject_grade_class_tables(), start=1):
        n_subj = si
        parts.extend(["", f"### 5.{si} {subj_full}"])
        for gi, (grade_full, rows) in enumerate(grade_blocks, start=1):
            parts.extend(["", f"#### 5.{si}.{gi} {grade_full}", "", md_table(rows)])
    if RESERVED:
        parts.extend(["", f"### 5.{n_subj + 1} 預留時段（不計小組班）"])
        for d, slot, room, teacher, title_t in RESERVED:
            parts.append(f"- {title_t}｜{teacher}｜{DAYS[d]} {SLOT_TIMES[slot]}｜{room}")
    parts.extend(
        [
            "",
            "## 6. 未排與待排",
            md_bullets(["本輪按老師已確定檔期與每周堂數編排", "不以「待確認老師」佔用課室格"]),
            "",
            "### 6.1 按老師",
            md_bullets(PENDING_BY_TEACHER),
            "",
            "### 6.2 按科目",
            md_bullets(PENDING_BY_SUBJECT),
            "",
            "## 7. 所有班別清單",
            md_bullets(
                [
                    "以下按星期與時段列出全部已排班及預留時段",
                    f"已排小組班合計 {len(CLASSES)} 班",
                    f"另預留時段 {len(RESERVED)} 個",
                ]
            ),
            "",
            md_table(weekly_summary_rows()),
            "",
            "## 附表 本輪老師回覆與出勤",
            md_bullets(
                [
                    "以下整理調查表各人意願，並結合各員工出勤、班數、科目",
                    "科目欄為營運確認，非表格 C 區自填（既有專科老師不經 C 區）",
                ]
            ),
            "",
            md_bullets(SURVEY_SLOT_NOTES),
            "",
            "### 附表.1 調查回覆",
            "",
            md_table(SURVEY_OVERVIEW),
            "",
            "### 附表.2 檔期與連堂約束",
            "",
            md_table(SURVEY_CONSTRAINTS),
            "",
            "### 附表.3 各員工出勤日、班數、科目",
            "",
            md_table(staff_appendix_table()),
        ]
    )
    parts.extend(["", "— 完 —", ""])
    return "\n".join(parts)


def build_weekly_md() -> str:
    parts: list[str] = [
        "# 明學教育 2627 學年常規專科班周時間表",
        "",
        f"獨立附件｜對應方案紀錄 **ver. {VERSION}**｜未定稿入庫",
        "",
        "> 本檔為獨立審閱稿，不置於方案正文。出 Word／PDF 時用 `--word`。",
        "",
        md_bullets(
            [
                f"已排小組班合計 {len(CLASSES)} 班；另預留時段 {len(RESERVED)} 個",
                "平日 15:15-16:30 或之前為返學時間，該列仍計空房",
                "最右欄為該時段空房數",
            ]
        ),
    ]
    for day_idx, day_name in enumerate(DAYS):
        parts.extend(
            [
                "",
                f"## {day_idx + 1}. {day_name}",
                "",
                md_table(day_grid(day_idx)),
            ]
        )
    parts.extend(["", "— 完 —", ""])
    return "\n".join(parts)


def build_teachers_md() -> str:
    parts: list[str] = [
        "# 明學教育 2627 學年各老師一周排程",
        "",
        f"獨立附件｜對應方案紀錄 **ver. {VERSION}**｜未定稿入庫",
        "",
        "> 本檔為審閱稿。出 Word／PDF 時用 `--word`。",
        "",
        md_bullets(
            [
                "本文件獨立於全校課室時間表，只按老師列出一周職務",
                "每位老師先列一周總覽，其後周視圖",
                "指定不排日於周視圖以「放假」或「不排」標示",
            ]
        ),
    ]
    for name, subject, n, days, note in STAFF:
        reserved_n = sum(1 for _d, _s, _r, teacher, _t in RESERVED if teacher == name)
        load = f"小組班：{n} 班"
        if reserved_n:
            load += f"；預留時段：{reserved_n} 個"
        parts.extend(
            [
                "",
                f"## {name}",
                md_bullets(
                    [
                        f"科目：{subject}",
                        load,
                        teacher_hours_text(n),
                        f"出勤日：{days}",
                        note,
                    ]
                ),
                "",
                "### 一周總覽",
                "",
                md_table(teacher_week_rows(name)),
                "",
                f"### {name}｜周視圖",
            ]
        )
        legend = week_view_legend(name)
        if legend:
            parts.extend(["", legend])
        grid_rows = [[cell for cell, _kind in row] for row in teacher_week_grid(name)]
        parts.extend(["", md_table(grid_rows)])
    parts.extend(["", "— 完 —", ""])
    return "\n".join(parts)


def write_markdown() -> tuple[Path, Path, Path]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    scheme_md = OUT_DIR / f"{STEM}.md"
    teachers_md = OUT_DIR / f"{STEM_TEACHERS}.md"
    weekly_md = OUT_DIR / f"{STEM_WEEKLY}.md"
    scheme_md.write_text(build_scheme_md(), encoding="utf-8")
    teachers_md.write_text(build_teachers_md(), encoding="utf-8")
    weekly_md.write_text(build_weekly_md(), encoding="utf-8")
    return scheme_md, teachers_md, weekly_md


def close_timetable_word_docs() -> None:
    script = '''
with timeout of 30 seconds
  tell application "Microsoft Word"
    repeat with d in (get documents)
      try
        if name of d contains "2627_timetable" then
          close d saving no
        end if
      end try
    end repeat
  end tell
end timeout
'''
    subprocess.run(["osascript", "-e", script], check=False, capture_output=True, text=True)


def refresh_word(docx_path: Path, pdf_path: Path) -> None:
    # Update Word TOC, then Save As PDF. Requires Microsoft Word on macOS.
    # Do not loop every field: TOC entries make that exceed AppleEvent timeout.
    script = f'''
with timeout of 1200 seconds
  tell application "Microsoft Word"
    set docPath to POSIX file "{docx_path}" as alias
    open docPath
    delay 3
    set theDoc to active document
    try
      if (count of tables of contents of theDoc) > 0 then
        update table of contents 1 of theDoc
      end if
    end try
    save theDoc
    set pdfFile to POSIX file "{pdf_path}" as string
    save as theDoc file name pdfFile file format format PDF
    delay 2
    try
      close active document saving no
    end try
  end tell
end timeout
'''
    result = subprocess.run(["osascript", "-e", script], check=False, capture_output=True, text=True)
    if result.returncode != 0:
        err = (result.stderr or result.stdout or "").strip() or f"exit {result.returncode}"
        raise RuntimeError(err)


def next_patch_version(version: str) -> str:
    major, minor = version.split(".")
    return f"{major}.{int(minor) + 1}"


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate 2627 timetable md (default) or Word/PDF")
    parser.add_argument(
        "--word",
        action="store_true",
        help="Also write docx and export PDF via Microsoft Word (built-in TOC / header / footer)",
    )
    args = parser.parse_args()
    validate()
    scheme_md, teachers_md, weekly_md = write_markdown()
    codes_csv = write_class_codes_csv()
    print("wrote", scheme_md)
    print("wrote", teachers_md)
    print("wrote", weekly_md)
    print("wrote", codes_csv)
    if not args.word:
        print("skip docx/pdf (pass --word when this version is ready to export)")
        return

    close_timetable_word_docs()
    docx_path = OUT_DIR / f"{STEM}.docx"
    pdf_path = OUT_DIR / f"{STEM}.pdf"
    build_docx(docx_path)
    print("wrote", docx_path)
    teachers_docx = OUT_DIR / f"{STEM_TEACHERS}.docx"
    teachers_pdf = OUT_DIR / f"{STEM_TEACHERS}.pdf"
    build_teachers_docx(teachers_docx)
    print("wrote", teachers_docx)
    weekly_docx = OUT_DIR / f"{STEM_WEEKLY}.docx"
    weekly_pdf = OUT_DIR / f"{STEM_WEEKLY}.pdf"
    build_weekly_docx(weekly_docx)
    print("wrote", weekly_docx)
    try:
        refresh_word(docx_path, pdf_path)
        print("wrote", pdf_path)
        refresh_word(teachers_docx, teachers_pdf)
        print("wrote", teachers_pdf)
        refresh_word(weekly_docx, weekly_pdf)
        print("wrote", weekly_pdf)
    except RuntimeError as exc:
        print("Word TOC / PDF update skipped:", exc)
        print("請在 Word 開啟 docx，更新目錄後另存 PDF")




if __name__ == "__main__":
    main()
