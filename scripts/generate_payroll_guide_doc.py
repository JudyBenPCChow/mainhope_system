#!/usr/bin/env python3
"""Generate payroll guide docx from markdown.

Source: docs/policies/staffing/PAYROLL_GUIDE.md
Output: docs/generated/payroll/PAYROLL_GUIDE.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "policies" / "staffing" / "PAYROLL_GUIDE.md"
OUTPUT = ROOT / "docs" / "generated" / "payroll" / "PAYROLL_GUIDE.docx"

FONT_NAME = "新細明體"
BODY_SIZE = 11
H1_SIZE = 14
H2_SIZE = 12
H3_SIZE = 11
TITLE_SIZE = 18
INK = RGBColor(0, 0, 0)


def set_run_font(run, size: float, *, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = False
    run.font.color.rgb = INK
    properties = run._element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    properties.append(color)
    language = OxmlElement("w:lang")
    language.set(qn("w:eastAsia"), "zh-HK")
    properties.append(language)


def set_style_font(style, size: float, *, bold: bool = False) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = INK
    properties = style.element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("**", "").replace("`", "").strip()


def add_runs_with_bold(paragraph, text: str, *, size: float = BODY_SIZE) -> None:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text).replace("`", "")
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if is_bold else part)
        set_run_font(run, size, bold=is_bold)


def add_bottom_rule(paragraph, *, size_eighths: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:pBdr"):
            p_pr.remove(child)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "000000")
    border.append(bottom)
    p_pr.append(border)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    element = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    element.extend([begin, instruction, separate, text, end])
    set_run_font(run, 10)


def add_footer_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        left = paragraph.add_run("— ")
        set_run_font(left, 10)
        add_page_field(paragraph)
        right = paragraph.add_run(" —")
        set_run_font(right, 10)


def add_para(document: Document, text: str, *, size: float = BODY_SIZE, space_after: float = 8) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    add_runs_with_bold(paragraph, text, size=size)


def add_heading(document: Document, text: str, *, level: int, page_break: bool = False) -> None:
    if page_break:
        document.add_page_break()
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 2"}[level]
    size = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}[level]
    paragraph = document.add_paragraph(strip_md_inline(text), style=style)
    paragraph.paragraph_format.space_before = Pt(0 if level == 1 else 14)
    paragraph.paragraph_format.space_after = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.line_spacing = 1.2
    for run in paragraph.runs:
        set_run_font(run, size, bold=True)
    if level == 1:
        add_bottom_rule(paragraph)


def add_bullet(document: Document, text: str, *, ordered: bool = False, index: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.4)
    prefix = f"{index}. " if ordered else "• "
    run = paragraph.add_run(prefix)
    set_run_font(run, BODY_SIZE)
    add_runs_with_bold(paragraph, text, size=BODY_SIZE)


def parse_table_row(line: str) -> list[str]:
    return [strip_md_inline(part) for part in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(set(cell) <= {"-", ":", " "} for cell in cells)


def set_cell_borders(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    for child in list(properties):
        if child.tag == qn("w:tcBorders"):
            properties.remove(child)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "000000")
        element.set(qn("w:space"), "0")
        borders.append(element)
    properties.append(borders)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        for column_index in range(columns):
            cell = table.rows[row_index].cells[column_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.2
            run = paragraph.add_run(row[column_index] if column_index < len(row) else "")
            set_run_font(run, 10, bold=(row_index == 0))
            set_cell_borders(cell)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    set_style_font(document.styles["Normal"], BODY_SIZE)
    set_style_font(document.styles["Heading 1"], H1_SIZE, bold=True)
    set_style_font(document.styles["Heading 2"], H2_SIZE, bold=True)


def build_document() -> Document:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)
    document.core_properties.title = "明學教育計糧指南"
    document.core_properties.author = "明學教育"

    index = 0
    if lines and lines[0].startswith("# "):
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(24)
        title.paragraph_format.space_after = Pt(12)
        run = title.add_run(strip_md_inline(lines[0][2:]))
        set_run_font(run, TITLE_SIZE, bold=True)
        add_bottom_rule(title, size_eighths=8)
        index = 1

    first_heading = True
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped.startswith("> "):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(re.sub(r"^>\s?", "", lines[index].strip()))
                index += 1
            add_para(document, "　".join(part for part in quote_lines if part), size=10)
            continue
        if stripped.startswith("## "):
            add_heading(document, stripped[3:], level=1, page_break=not first_heading)
            first_heading = False
            index += 1
            continue
        if stripped.startswith("### "):
            add_heading(document, stripped[4:], level=2)
            index += 1
            continue
        if stripped.startswith("#### "):
            add_heading(document, stripped[5:], level=3)
            index += 1
            continue
        if stripped.startswith("|"):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                row = parse_table_row(lines[index])
                if not is_separator_row(row):
                    rows.append(row)
                index += 1
            add_table(document, rows)
            continue
        if stripped.startswith("```"):
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_para(document, "\n".join(code_lines), size=10.5, space_after=10)
            index += 1
            continue
        ordered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ordered:
            add_bullet(document, ordered.group(2), ordered=True, index=int(ordered.group(1)))
            index += 1
            continue
        if stripped.startswith("- "):
            add_bullet(document, stripped[2:])
            index += 1
            continue
        add_para(document, stripped)
        index += 1

    add_footer_page_numbers(document)
    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    main()
